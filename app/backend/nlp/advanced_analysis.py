from typing import Dict, List, Tuple, Any, Set, Optional
import logging
import os
import re
import spacy
from datetime import datetime
from .__init__ import NLPAnalyzer
from .summary_extraction import extract_summary, clean_sensitive_information
from .skill_keywords import (
    PROGRAMMING_LANGUAGES, WEB_TECHNOLOGIES, DATA_SCIENCE,
    CLOUD_DEVOPS, DATABASE_TECHNOLOGIES, SOFT_SKILLS, BUSINESS_SKILLS,
    ALL_SKILLS, SKILL_CATEGORIES
)
# Import specialized extraction functions for complex resume formats
try:
    from .specialized_extraction import (
        extract_spaced_name, extract_contact_info_complex, extract_section_with_mixed_layout,
        extract_education_complex, extract_experience_complex, extract_skills_complex,
        extract_summary_complex, detect_complex_resume
    )
    SPECIALIZED_EXTRACTION_AVAILABLE = True
except ImportError:
    logging.warning("Specialized extraction functions not available. Complex resume formats may not be parsed correctly.")
    SPECIALIZED_EXTRACTION_AVAILABLE = False

# Import document parsing libraries with error handling
try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    import docx
except ImportError:
    docx = None

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

from .skill_keywords import ALL_SKILLS, SKILL_CATEGORIES, ALL_JOB_TITLES, JOB_TITLES, PROGRAMMING_LANGUAGES, WEB_TECHNOLOGIES, DATA_SCIENCE, CLOUD_DEVOPS, BUSINESS_SKILLS, SOFT_SKILLS

def parse_cv(filepath: str) -> dict:
    """
    Advanced CV analysis: Parse a CV (PDF or DOCX) and extract structured profile info, section-by-section analysis, skills, ATS/bias checks, and actionable recommendations.
    Returns: dict with keys: name, target_job, skills, education, experience, email, sections, ats_report, bias_report, language_report, summary, recommendations.
    """
    import textblob
    from datetime import datetime

    try:
        # Extract text from document based on file type
        text = extract_text_from_document(filepath)

        # Load NLP model
        try:
            nlp = spacy.load("en_core_web_sm")
        except Exception as e:
            logging.error(f"Error loading spaCy model: {e}")
            # Fallback to a simpler approach if spaCy fails
            return create_minimal_profile_from_text(text)

        # Process text with spaCy
        doc = nlp(text[:1000000])  # Limit to 1M chars to prevent memory issues

        # Check if this is a complex resume format that needs specialized extraction
        is_complex_format = False
        if SPECIALIZED_EXTRACTION_AVAILABLE:
            is_complex_format = detect_complex_resume(text)
            if is_complex_format:
                logging.info("Detected complex resume format. Using specialized extraction functions.")

        # Extract name with appropriate method
        if is_complex_format:
            # Try specialized extraction for spaced names (e.g., "M O H A M M E D")
            spaced_name = extract_spaced_name(text)
            if spaced_name:
                name = spaced_name
            else:
                name = extract_name(text, doc)
        else:
            name = extract_name(text, doc)

        # Extract contact information
        if is_complex_format:
            contact_info = extract_contact_info_complex(text)
        else:
            contact_info = extract_contact_info(text)

        email = contact_info.get('email', '')
        phone = contact_info.get('phone', '')
        location = contact_info.get('location', '')

        # Try direct parsing for common CV format with ALL CAPS headers
        sections = {}
        lines = text.split('\n')

        # Direct mapping for ALL CAPS headers (common in plain text CVs)
        direct_mapping = {
            "SUMMARY": "summary",
            "PROFILE": "summary",
            "OBJECTIVE": "summary",
            "ABOUT ME": "summary",
            "EXPERIENCE": "experience",
            "WORK EXPERIENCE": "experience",
            "PROFESSIONAL EXPERIENCE": "experience",
            "EMPLOYMENT HISTORY": "experience",
            "EDUCATION": "education",
            "ACADEMIC BACKGROUND": "education",
            "EDUCATIONAL BACKGROUND": "education",
            "SKILLS": "skills",
            "TECHNICAL SKILLS": "skills",
            "CORE COMPETENCIES": "skills",
            "PROJECTS": "projects",
            "PROJECT EXPERIENCE": "projects",
            "CERTIFICATIONS": "certifications",
            "CERTIFICATES": "certifications",
            "AWARDS": "awards",
            "HONORS": "awards",
            "LANGUAGES": "languages",
            "LANGUAGE PROFICIENCY": "languages",
            "INTERESTS": "interests",
            "HOBBIES": "interests",
            "REFERENCES": "references"
        }

        # Log the first few lines for debugging
        logging.info(f"First 10 lines of CV text:")
        for i, line in enumerate(lines[:10]):
            logging.info(f"Line {i}: '{line}'")

        # Find all uppercase section headers
        section_indices = []
        for i, line in enumerate(lines):
            line = line.strip()
            # Check for ALL CAPS section headers
            if line and line.isupper() and len(line) >= 4 and len(line) <= 30:
                logging.info(f"Found potential header: '{line}'")
                if line in direct_mapping:
                    section_indices.append((i, direct_mapping[line], line))
                    logging.info(f"Found section header: '{line}' -> {direct_mapping[line]}")
            # Also check for section headers that are not all caps but match our mapping
            elif line and len(line) >= 4 and len(line) <= 30:
                for header, section in direct_mapping.items():
                    if line.upper() == header:
                        section_indices.append((i, section, line))
                        logging.info(f"Found section header (case insensitive): '{line}' -> {section}")
                        break

        # Extract content between section headers
        if len(section_indices) >= 2:
            for i in range(len(section_indices)):
                start_idx = section_indices[i][0] + 1  # Start after the header
                end_idx = section_indices[i+1][0] if i < len(section_indices) - 1 else len(lines)
                section_name = section_indices[i][1]

                # Extract content
                content = '\n'.join(lines[start_idx:end_idx]).strip()
                if content:
                    sections[section_name] = content
                    logging.info(f"Extracted {section_name} section with {len(content)} characters")

        # If direct parsing didn't work, fall back to the more complex extraction
        if len(sections) < 2:
            sections = extract_sections(text)

        # Get common section content
        experience_text = get_section_content(sections, ["experience", "work experience", "professional experience", "employment history", "work history"])
        education_text = get_section_content(sections, ["education", "academic background", "educational background", "academic history", "qualifications"])
        skills_section = get_section_content(sections, ["skills", "technical skills", "core competencies", "key skills", "expertise"])
        projects = get_section_content(sections, ["projects", "project experience", "key projects", "relevant projects"])

        # Log the section content for debugging
        logging.info(f"Experience section: {experience_text[:100] if experience_text else 'Not found'}")
        logging.info(f"Education section: {education_text[:100] if education_text else 'Not found'}")
        logging.info(f"Skills section: {skills_section[:100] if skills_section else 'Not found'}")

        # Extract structured education and experience information
        if is_complex_format and SPECIALIZED_EXTRACTION_AVAILABLE:
            # Use specialized extraction for complex resume formats
            education = extract_education_complex(text)
            if not education and education_text:
                # Fall back to standard extraction if specialized extraction fails
                education = extract_education(education_text, doc)

            experience = extract_experience_complex(text)
            if not experience and experience_text:
                # Fall back to standard extraction if specialized extraction fails
                experience = extract_experience(experience_text, doc)

            # Extract skills with specialized detection
            skills = extract_skills_complex(text)
            if not skills:
                # Fall back to standard extraction if specialized extraction fails
                skills = extract_skills(text, skills_section, doc)
        else:
            # Use standard extraction for normal resume formats
            education = extract_education(education_text, doc)
            experience = extract_experience(experience_text, doc)
            skills = extract_skills(text, skills_section, doc)

        # Extract or generate a professional summary
        if is_complex_format and SPECIALIZED_EXTRACTION_AVAILABLE:
            # Use specialized extraction for complex resume formats
            summary = extract_summary_complex(text)
            if not summary:
                # Fall back to standard extraction if specialized extraction fails
                summary = extract_summary(text, sections, doc)
        else:
            # Use standard extraction for normal resume formats
            summary = extract_summary(text, sections, doc)

        logging.info(f"Summary: {summary[:100] if summary else 'Not found'}")

        # Extract target job
        target_job = extract_target_job(text, summary, doc)

        # Score sections
        section_scores = {
            "summary": score_section(summary, skills),
            "experience": score_section(experience_text, skills),
            "education": score_section(education_text, skills),
            "skills": score_section(skills_section, skills),
            "projects": score_section(projects, skills)
        }

        # Generate language quality report
        language_report = analyze_language_quality(text)

        # Generate ATS report
        ats_report = generate_ats_report(text, email, phone, skills, experience, education, skills_section, target_job)

        # Generate bias report
        bias_report = check_for_bias(text)

        # Generate recommendations
        recommendations = generate_recommendations(section_scores, ats_report, bias_report, skills, target_job)

        # Print sections for debugging
        logging.info(f"Extracted sections: {list(sections.keys())}")
        if 'education' in sections:
            logging.info(f"Education section content: {sections['education'][:100]}...")
        if 'experience' in sections:
            logging.info(f"Experience section content: {sections['experience'][:100]}...")

        # Clean the summary one more time to ensure no sensitive information is included
        if summary:
            # Remove any remaining phone numbers or email addresses that might have been missed
            summary = clean_sensitive_information(summary)

            # Make sure the summary doesn't contain the raw phone number
            if phone and phone in summary:
                summary = summary.replace(phone, '')

            # Make sure the summary doesn't contain the email address
            if email and email in summary:
                summary = summary.replace(email, '')

            # Clean up any artifacts from the removal
            summary = re.sub(r'\s+', ' ', summary).strip()
            summary = re.sub(r'^[,.\s:;-]+|[,.\s:;-]+$', '', summary).strip()

        # Return complete profile
        return {
            'name': name,
            'email': email,
            'phone': phone,
            'location': location,
            'skills': skills,
            'education': education,
            'experience': experience,
            'target_job': target_job,
            'sections': sections,
            'section_scores': section_scores,
            'ats_report': ats_report,
            'bias_report': bias_report,
            'language_report': language_report,
            'summary': summary,
            'recommendations': recommendations,
            'raw_text': text[:5000],  # Include limited raw text for debugging
            'uploaded': True,
            'lastUpdated': datetime.now().isoformat(),
            'complex_format_detected': is_complex_format  # Add flag to indicate if complex format was detected
        }
    except Exception as e:
        import traceback
        logging.error(f"Error in parse_cv: {e}")
        logging.error(f"Traceback: {traceback.format_exc()}")
        # If advanced parsing fails, fall back to minimal profile
        try:
            # Extract text from document
            text = extract_text_from_document(filepath)
            return create_minimal_profile_from_text(text)
        except Exception as e2:
            logging.error(f"Error in fallback to create_minimal_profile: {e2}")
            logging.error(f"Fallback traceback: {traceback.format_exc()}")
            # Return a very minimal profile if everything fails
            return {
                'name': 'User',
                'email': '',
                'phone': '',
                'location': '',
                'skills': [],
                'education': [],
                'experience': [],
                'summary': '',
                'uploaded': True,
                'lastUpdated': datetime.now().isoformat(),
                'complex_format_detected': False
            }

def extract_text_from_document(filepath: str) -> str:
    """
    Extract text from various document formats with enhanced error handling and format support.
    Supports PDF, DOCX, DOC, TXT, RTF, and HTML files.
    Uses multiple extraction methods for each format to maximize success rate.
    """
    text = ""
    ext = os.path.splitext(filepath)[-1].lower()

    try:
        # PDF extraction with multiple fallbacks
        if ext == ".pdf":
            pdf_text = ""
            extraction_methods = []

            # Method 1: pdfplumber
            if pdfplumber is not None:
                extraction_methods.append(("pdfplumber", lambda: extract_with_pdfplumber(filepath)))

            # Method 2: PyPDF2
            if PyPDF2 is not None:
                extraction_methods.append(("PyPDF2", lambda: extract_with_pypdf2(filepath)))

            # Try each method until one succeeds
            for method_name, extract_func in extraction_methods:
                try:
                    pdf_text = extract_func()
                    if pdf_text.strip():
                        logging.info(f"Successfully extracted PDF text using {method_name}")
                        text = pdf_text
                        break
                except Exception as e:
                    logging.warning(f"{method_name} extraction failed: {e}")

            # If all methods failed, try a basic approach
            if not text.strip():
                logging.warning("All PDF extraction methods failed, trying basic approach")
                try:
                    with open(filepath, 'rb') as f:
                        content = f.read()
                        # Simple text extraction from PDF binary
                        text = ' '.join(str(content).replace('\\n', ' ').split())
                except Exception as e:
                    logging.error(f"Basic PDF extraction failed: {e}")

        # Word document extraction
        elif ext in (".docx", ".doc"):
            if docx is not None:
                try:
                    doc = docx.Document(filepath)
                    # Extract text from paragraphs
                    paragraphs = [p.text for p in doc.paragraphs]

                    # Extract text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                paragraphs.append(cell.text)

                    text = "\n".join(paragraphs)
                except Exception as e:
                    logging.error(f"Error extracting text from DOCX: {e}")
                    # Try a more basic approach for .doc files
                    if ext == ".doc":
                        try:
                            # Try to import textract, but don't fail if it's not available
                            try:
                                import textract
                                text = textract.process(filepath).decode('utf-8')
                            except ImportError:
                                logging.warning("Textract not available, using basic text extraction")
                                # Basic approach - try to read as binary and extract text
                                with open(filepath, 'rb') as f:
                                    content = f.read()
                                    # Simple text extraction from binary
                                    text = ' '.join(str(content).replace('\\n', ' ').split())
                        except Exception as e2:
                            logging.error(f"Basic DOC extraction failed: {e2}")
                            raise ValueError(f"Could not extract text from DOC file: {e}")
                    else:
                        raise ValueError(f"Could not extract text from DOCX file: {e}")
            else:
                raise ValueError("Python-docx library not available")

        # Plain text files
        elif ext == ".txt":
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()

        # RTF files
        elif ext == ".rtf":
            # Basic RTF parsing that doesn't require additional libraries
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                rtf_text = f.read()
                # Basic RTF cleaning
                text = re.sub(r'\\[a-z]+', ' ', rtf_text)
                text = re.sub(r'\{|\}|\\', '', text)

        # HTML files
        elif ext in (".html", ".htm"):
            # Basic HTML parsing that doesn't require additional libraries
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                html = f.read()
                # Basic HTML tag removal
                text = re.sub(r'<[^>]+>', ' ', html)
                text = re.sub(r'\s+', ' ', text).strip()

        else:
            logging.warning(f"Unsupported file format: {ext}, attempting basic text extraction")
            # Try basic text extraction for unknown formats
            try:
                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
            except Exception as e:
                logging.error(f"Basic text extraction failed: {e}")
                raise ValueError(f"Unsupported file format: {ext}")

    except Exception as e:
        logging.error(f"Error extracting text from {filepath}: {e}")
        raise ValueError(f"Could not extract text from file: {e}")

    # Post-processing to clean up the text
    if text:
        # Fix common encoding issues
        text = text.replace('\x00', '')

        # For text files, preserve the original line breaks
        if ext == ".txt":
            # Just normalize whitespace within lines
            lines = text.split('\n')
            text = '\n'.join(re.sub(r'\s+', ' ', line).strip() for line in lines if line.strip())
        else:
            # For other formats, normalize whitespace and then try to detect section breaks
            # First normalize all whitespace
            text = re.sub(r'\s+', ' ', text)

            # Then try to detect and restore section breaks
            # Look for patterns like "SECTION HEADER" (all caps words)
            text = re.sub(r'([A-Z]{4,})', r'\n\1', text)

            # Look for patterns like "Name: John Smith" (label followed by colon)
            text = re.sub(r'([A-Za-z]+:)', r'\n\1', text)

            # Split into lines for better readability
            text = '\n'.join(line.strip() for line in text.split('\n') if line.strip())

    return text

def extract_with_pdfplumber(filepath: str) -> str:
    """Extract text from PDF using pdfplumber."""
    if pdfplumber is None:
        raise ImportError("pdfplumber is not installed")

    with pdfplumber.open(filepath) as pdf:
        pages = []
        for page in pdf.pages:
            try:
                page_text = page.extract_text() or ''
                pages.append(page_text)
            except Exception as e:
                logging.warning(f"Error extracting text from page: {e}")
                pages.append("")
        return "\n".join(pages)

def extract_with_pypdf2(filepath: str) -> str:
    """Extract text from PDF using PyPDF2."""
    if PyPDF2 is None:
        raise ImportError("PyPDF2 is not installed")

    with open(filepath, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        pages = []
        for page in reader.pages:
            try:
                page_text = page.extract_text() or ''
                pages.append(page_text)
            except Exception as e:
                logging.warning(f"Error extracting text from page: {e}")
                pages.append("")
        return "\n".join(pages)

def create_minimal_profile_from_text(text: str) -> Dict[str, Any]:
    """Create a minimal profile from text when advanced parsing fails."""
    try:
        # Extract basic information
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        email_matches = re.findall(email_pattern, text)
        email = email_matches[0] if email_matches else ""

        # Extract phone number
        phone_pattern = r'(?:\+\d{1,3}[-.\s]?)?(?:\(?\d{1,4}\)?[-.\s]?)?(?:\d{1,4}[-.\s]?){1,3}\d{1,4}'
        phone_matches = re.findall(phone_pattern, text)
        phone = phone_matches[0] if phone_matches else ""

        # Return a very minimal profile
        return {
            'name': 'User',
            'email': email,
            'phone': phone,
            'location': '',
            'skills': [],
            'education': [],
            'experience': [],
            'summary': '',
            'uploaded': True,
            'lastUpdated': datetime.now().isoformat(),
            'complex_format_detected': False
        }
    except Exception as e:
        logging.error(f"Error in create_minimal_profile_from_text: {e}")
        # Return a very minimal profile if everything fails
        return {
            'name': 'User',
            'email': '',
            'phone': '',
            'location': '',
            'skills': [],
            'education': [],
            'experience': [],
            'summary': '',
            'uploaded': True,
            'lastUpdated': datetime.now().isoformat(),
            'complex_format_detected': False
        }

# Alias for backward compatibility
def create_minimal_profile(filepath: str) -> Dict[str, Any]:
    """Create a minimal profile from a file path."""
    try:
        text = extract_text_from_document(filepath)
        return create_minimal_profile_from_text(text)
    except Exception as e:
        logging.error(f"Error in create_minimal_profile: {e}")
        return {
            'name': 'User',
            'email': '',
            'phone': '',
            'location': '',
            'skills': [],
            'education': [],
            'experience': [],
            'summary': '',
            'uploaded': True,
            'lastUpdated': datetime.now().isoformat(),
            'complex_format_detected': False
        }

def extract_name(text: str, doc) -> str:
    """
    Enhanced name extraction using multiple strategies.

    1. Look for names at the beginning of the document
    2. Use NER to identify PERSON entities
    3. Look for patterns that typically indicate names
    4. Apply validation rules to filter out false positives
    5. Use contact information context to validate names
    6. Apply cultural name pattern recognition

    Note: This function has been improved to avoid the "Sarah" issue where
    the system would incorrectly assume a default name.
    """
    # Company/organization keywords to filter out
    company_keywords = [
        "inc", "llc", "university", "college", "ltd", "corp", "school", "institute",
        "gmbh", "co.", "company", "group", "plc", "associates", "partners", "corporation",
        "foundation", "committee", "society", "organization", "centre", "center",
        "technologies", "solutions", "systems", "international", "global", "worldwide",
        "consulting", "services", "academy", "prep", "ai", "khaimah", "resume", "cv",
        "curriculum vitae", "application", "profile", "portfolio", "limited", "holdings",
        "enterprises", "industries", "agency", "bureau", "department", "division", "office"
    ]

    # Common names to avoid using as defaults (to prevent the "Sarah" issue)
    common_default_names = [
        "sarah", "john", "jane", "bob", "alice", "david", "michael", "james",
        "mary", "robert", "jennifer", "linda", "william", "richard", "susan",
        "joseph", "thomas", "charles", "daniel", "matthew", "anthony", "donald",
        "mark", "paul", "steven", "andrew", "kenneth", "george", "joshua", "kevin",
        "brian", "edward", "ronald", "timothy", "jason", "jeffrey", "ryan", "jacob",
        "gary", "nicholas", "eric", "stephen", "jonathan", "larry", "justin", "scott",
        "brandon", "benjamin", "samuel", "gregory", "alexander", "patrick", "frank",
        "raymond", "jack", "dennis", "jerry", "tyler", "aaron", "jose", "adam", "nathan",
        "henry", "douglas", "zachary", "peter", "kyle", "walter", "harold", "jeremy",
        "ethan", "carl", "keith", "roger", "gerald", "christian", "terry", "sean", "arthur",
        "austin", "noah", "lawrence", "jesse", "joe", "bryan", "billy", "jordan", "albert",
        "dylan", "bruce", "willie", "gabriel", "alan", "juan", "logan", "wayne", "ralph",
        "roy", "eugene", "randy", "vincent", "russell", "louis", "philip", "bobby", "johnny",
        "bradley", "elizabeth", "patricia", "barbara", "jessica", "margaret", "ashley",
        "nancy", "lisa", "betty", "dorothy", "sandra", "ashley", "kimberly", "donna",
        "emily", "michelle", "carol", "amanda", "melissa", "deborah", "stephanie", "rebecca",
        "laura", "sharon", "cynthia", "kathleen", "amy", "shirley", "angela", "helen",
        "anna", "brenda", "pamela", "nicole", "emma", "samantha", "katherine", "christine",
        "debra", "rachel", "catherine", "carolyn", "janet", "ruth", "maria", "heather",
        "diane", "virginia", "julie", "joyce", "victoria", "olivia", "kelly", "christina",
        "lauren", "joan", "evelyn", "judith", "megan", "cheryl", "hannah", "martha",
        "andrea", "frances", "denise", "ruby", "grace", "beverly", "theresa", "judy",
        "madison", "sophia", "marie", "gloria", "doris", "sara", "janice", "julia",
        "abigail", "ann", "kathryn", "jacqueline", "jean"
    ]

    # Common CV section headers to filter out
    section_headers = [
        "summary", "profile", "objective", "experience", "education", "skills", "projects",
        "certifications", "awards", "publications", "languages", "interests", "references",
        "contact", "personal", "professional", "qualifications", "achievements", "expertise",
        "competencies", "capabilities", "strengths", "employment", "history", "background",
        "career", "objective", "goal", "mission", "vision", "statement", "highlights"
    ]

    # Common English words that might be mistaken for names
    common_words = [
        "strong", "resume", "curriculum", "vitae", "professional", "career", "job",
        "application", "position", "qualified", "experienced", "skilled", "expertise",
        "proficient", "competent", "accomplished", "successful", "effective", "efficient",
        "reliable", "dependable", "trustworthy", "honest", "integrity", "dedicated",
        "committed", "motivated", "enthusiastic", "passionate", "driven", "determined",
        "ambitious", "proactive", "innovative", "creative", "analytical", "detail",
        "oriented", "organized", "methodical", "systematic", "thorough", "comprehensive",
        "adaptable", "flexible", "versatile", "dynamic", "energetic", "resourceful",
        "contact", "information", "address", "phone", "email", "website", "linkedin",
        "github", "portfolio", "reference", "recommendation", "confidential", "private"
    ]

    def is_company_or_institution(text_to_check: str) -> bool:
        """More robust check for company/organization names"""
        text_lower = text_to_check.lower()

        # Check for company keywords
        if any(kw in text_lower.split() for kw in company_keywords):
            return True

        # Check for typical company patterns (e.g., Capital letters with spaces)
        if re.match(r'^[A-Z][a-z]+ [A-Z][a-z]+( [A-Z][a-z]+)*$', text_to_check) and len(text_to_check.split()) >= 3:
            return True

        # Check for all caps (often organizations)
        if text_to_check.isupper() and len(text_to_check) > 3:
            return True

        # Check for company-like patterns (e.g., "ABC Corp", "XYZ Inc")
        if re.match(r'^[A-Z]{2,}(\s+[A-Za-z]+){1,2}$', text_to_check):
            return True

        return False

    def is_valid_name(name_to_check: str) -> bool:
        """Check if a string looks like a valid person name"""
        # Remove non-alphabetic characters except spaces, hyphens, and apostrophes
        clean_name = re.sub(r'[^a-zA-Z\s\'-]', '', name_to_check).strip()

        # Too short or too long
        if len(clean_name) < 3 or len(clean_name) > 40:
            return False

        # Too many words
        if len(clean_name.split()) > 4:
            return False

        # Check if it's a company/institution
        if is_company_or_institution(clean_name):
            return False

        # Check if it contains common CV section headers
        if any(header in clean_name.lower() for header in section_headers):
            return False

        # If the name is a single word and it's in our list of common words, it's probably not a name
        if len(clean_name.split()) == 1 and clean_name.lower() in common_words:
            return False

        # Check if the name is a common default name (to prevent the "Sarah" issue)
        # Only reject if it's a single common name without a last name
        if len(clean_name.split()) == 1 and clean_name.lower() in common_default_names:
            logging.warning(f"Rejected common default name: {clean_name}")
            return False

        # Check if all words are properly capitalized (typical for names)
        # Allow for hyphenated names (e.g., "Jean-Pierre")
        words = []
        for word in clean_name.split():
            words.extend([w for w in re.split(r'[-]', word) if w])

        if not all(w[0].isupper() for w in words if w and len(w) > 1):
            return False

        return True

    def extract_contact_context(text: str, lines: list) -> dict:
        """Extract contact information context to help validate names"""
        context = {}

        # Look for email addresses
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        if emails:
            context['email'] = emails[0]
            # Try to extract name from email
            email_name_match = re.match(r'([a-zA-Z]+)\.?([a-zA-Z]+)?@', emails[0])
            if email_name_match:
                if email_name_match.group(2):  # Has both first and last name
                    context['email_name'] = f"{email_name_match.group(1).capitalize()} {email_name_match.group(2).capitalize()}"
                else:  # Only has one name component
                    context['email_name'] = email_name_match.group(1).capitalize()

        # Look for LinkedIn profiles
        linkedin_pattern = r'linkedin\.com/in/([a-zA-Z0-9_-]+)'
        linkedin_matches = re.findall(linkedin_pattern, text.lower())
        if linkedin_matches:
            linkedin_name = linkedin_matches[0].replace('-', ' ').replace('_', ' ')
            context['linkedin_name'] = ' '.join(word.capitalize() for word in linkedin_name.split())

        return context

    # Strategy 1: Look for names in the first few lines (common CV format)
    name_candidates = []

    # Check first 7 non-empty lines (increased from 5)
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    for i, line in enumerate(lines[:7]):
        # Only consider lines that look like names (1-3 words, properly capitalized)
        if is_valid_name(line) and len(line.split()) <= 3:
            # Higher confidence for lines at the very beginning
            position_boost = max(0, (7 - i) / 10)  # 0.7 for first line, 0.6 for second, etc.

            # Single word names need higher scrutiny - they're often headers or titles
            if len(line.split()) == 1:
                # Lower confidence for single-word names
                name_candidates.append((line, 0.6 + position_boost))
            else:
                # Higher confidence for multi-word names (e.g., "John Smith")
                name_candidates.append((line, 0.8 + position_boost))

    # Strategy 2: Use NER to identify PERSON entities
    person_entities = [ent.text.strip() for ent in doc.ents if ent.label_ == "PERSON"]
    for entity in person_entities:
        if is_valid_name(entity):
            # Higher confidence for entities near the beginning
            first_occurrence = text.find(entity)
            position_score = max(0.1, 1.0 - (first_occurrence / 1000))
            name_candidates.append((entity, 0.7 + position_score))

    # Strategy 3: Look for patterns that typically indicate names
    name_patterns = [
        # Common name label patterns with high confidence
        (r'(?i)(?:full\s+)?name\s*[:;-]\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+){0,3})', 0.95),
        (r'(?i)candidate\s*[:;-]\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+){0,3})', 0.9),
        (r'(?i)applicant\s*[:;-]\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+){0,3})', 0.9),
        (r'(?i)contact\s*[:;-](?:[^:;-]*?[^A-Za-z])?([A-Z][a-z]+(?:\s+[A-Z][a-z]+){0,3})', 0.85),
        # Cultural name patterns (e.g., "Firstname M. Lastname")
        (r'(?i)([A-Z][a-z]+\s+[A-Z]\.?\s+[A-Z][a-z]+)', 0.9),
        # Generic patterns with medium confidence
        (r'(?i)^([A-Z][a-z]+(?:\s+[A-Z][a-z]+){0,3})$', 0.6)
    ]

    for pattern, confidence in name_patterns:
        matches = re.findall(pattern, text)
        for match in matches:
            if is_valid_name(match):
                # Higher confidence for matches with explicit name labels
                name_candidates.append((match, confidence))

    # Strategy 4: Use contact information context to validate names
    contact_context = extract_contact_context(text, lines)

    # If we have an email-derived name, add it as a candidate
    if 'email_name' in contact_context and is_valid_name(contact_context['email_name']):
        name_candidates.append((contact_context['email_name'], 0.75))

    # If we have a LinkedIn-derived name, add it as a candidate
    if 'linkedin_name' in contact_context and is_valid_name(contact_context['linkedin_name']):
        name_candidates.append((contact_context['linkedin_name'], 0.7))

    # Strategy 5: Look for names in context with contact information
    if 'email' in contact_context:
        email = contact_context['email']
        # Find the line containing the email
        for i, line in enumerate(lines):
            if email in line:
                # Check the line before and after for potential names
                for j in [i-1, i+1]:
                    if 0 <= j < len(lines) and is_valid_name(lines[j]) and len(lines[j].split()) <= 3:
                        name_candidates.append((lines[j], 0.8))
                break

    # Sort candidates by confidence
    name_candidates.sort(key=lambda x: x[1], reverse=True)

    # Deduplicate candidates (case-insensitive)
    seen = set()
    unique_candidates = []
    for candidate, confidence in name_candidates:
        if candidate.lower() not in seen:
            seen.add(candidate.lower())
            unique_candidates.append((candidate, confidence))
    name_candidates = unique_candidates

    # Return the highest confidence candidate if it meets our threshold
    if name_candidates and name_candidates[0][1] >= 0.8:
        candidate_name = name_candidates[0][0]
        # Final check to avoid returning a common default name
        if len(candidate_name.split()) == 1 and candidate_name.lower() in common_default_names:
            logging.warning(f"Avoided returning common default name: {candidate_name}")
            return ""  # Return empty string instead of a common default name
        return candidate_name

    # If we have candidates but none with high confidence, check if any are multi-word names
    # Multi-word names are more likely to be actual names than single words
    for candidate, confidence in name_candidates:
        if len(candidate.split()) >= 2:
            # Final check to avoid returning a name that starts with a common default name
            first_name = candidate.split()[0]
            if first_name.lower() in common_default_names and len(candidate.split()) < 3:
                # Only accept if it has at least 3 parts (first, middle, last) when first name is common
                logging.warning(f"Avoided returning name with common first name: {candidate}")
                continue
            return candidate

    # If we still don't have a good candidate, return empty string
    # The frontend will handle this by showing a placeholder or asking the user
    logging.info("No suitable name found in CV, returning empty string")
    return ""

def extract_contact_info(text: str) -> Dict[str, str]:
    """Extract contact information (email, phone) from text with enhanced phone detection."""
    contact_info = {}

    # Extract email using regex
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    emails = re.findall(email_pattern, text)
    if emails:
        contact_info['email'] = emails[0]

    # Extract phone numbers using comprehensive regex patterns
    # This covers various international phone number formats
    phone_patterns = [
        # North American formats
        r'\b(?:\+?1[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b',  # (123) 456-7890, 123-456-7890, +1 123-456-7890

        # International formats with country codes
        r'\b\+\d{1,3}[-.\s]?\d{1,14}\b',  # +XX XXXXXXXXXX (any international format)

        # Common formats without separators
        r'\b\d{10,12}\b',  # 1234567890, 123456789012

        # Short formats
        r'\b\d{3}[-.\s]?\d{4}\b',  # 123-4567

        # UAE specific formats
        r'\b(?:\+?971|0)[-.\s]?(?:50|55|56|58|2|3|4|6|7|9)\d{7}\b',  # +971 50 1234567, 050 1234567

        # Format with "Tel:" or "Phone:" prefix
        r'(?:Tel|Phone|Mobile|Cell|Contact)(?::|number|#|is|at)?[-.\s]*(?:\+?\d[-.\s\d]{8,})'
    ]

    # Look for phone numbers with context first (higher confidence)
    context_patterns = [
        r'(?:phone|mobile|cell|tel|telephone|contact)(?:\s*(?:number|#|:|\())?[^\n\r]*?(\+?[\d\s\(\)\-\.]{7,})',
        r'(?:phone|mobile|cell|tel|telephone|contact)[^\n\r]*?(\+?[\d\s\(\)\-\.]{7,})',
        r'(\+?[\d\s\(\)\-\.]{7,})[^\n\r]*?(?:phone|mobile|cell|tel|telephone|contact)'
    ]

    # Try context patterns first (more likely to be actual phone numbers)
    for pattern in context_patterns:
        matches = re.findall(pattern, text.lower())
        if matches:
            # Clean up the match to get just the digits and common separators
            phone = re.sub(r'[^\d\+\-\.\(\)\s]', '', matches[0]).strip()
            if phone and len(re.sub(r'[^\d]', '', phone)) >= 7:  # At least 7 digits
                contact_info['phone'] = phone
                return contact_info

    # If no context-based phone found, try the regular patterns
    for pattern in phone_patterns:
        phones = re.findall(pattern, text)
        if phones:
            # Take the first match but clean it up
            phone = re.sub(r'[^\d\+\-\.\(\)\s]', '', phones[0]).strip()
            if phone and len(re.sub(r'[^\d]', '', phone)) >= 7:  # At least 7 digits
                contact_info['phone'] = phone
                break

    return contact_info

def extract_sections(text: str) -> Dict[str, str]:
    """
    Enhanced section detection with fuzzy matching and better boundary detection.
    Returns a dictionary of section name -> section content.

    This improved version uses multiple strategies to accurately identify sections:
    1. Looks for ALL CAPS headers (common in plain text CVs)
    2. Detects headers with specific formatting (bold, underlined in PDFs)
    3. Uses regex patterns to identify section boundaries
    4. Employs fuzzy matching for section names
    5. Handles nested sections and subsections
    """
    # Common section headers with variations - expanded for better matching
    section_headers = {
        "summary": ["summary", "professional summary", "profile", "professional profile", "about me", "objective",
                   "career objective", "personal statement", "personal profile", "career summary", "executive summary",
                   "overview", "professional overview", "career profile", "personal summary", "introduction"],
        "experience": ["experience", "work experience", "professional experience", "employment history", "work history",
                      "career history", "professional background", "professional history", "employment experience",
                      "work background", "career experience", "job history", "positions held", "professional positions",
                      "employment", "work", "career", "jobs", "professional career"],
        "education": ["education", "educational background", "academic background", "qualifications", "academic qualifications",
                     "educational qualifications", "academic history", "academic record", "educational history",
                     "academic achievements", "academic experience", "educational experience", "academic credentials",
                     "academic profile", "educational profile", "schooling", "academic training", "formal education"],
        "skills": ["skills", "technical skills", "core competencies", "key skills", "expertise", "proficiencies",
                  "capabilities", "skill set", "areas of expertise", "areas of knowledge", "competencies",
                  "professional skills", "key competencies", "strengths", "qualifications", "technical expertise",
                  "professional competencies", "key capabilities", "core skills", "specialties", "specializations"],
        "projects": ["projects", "project experience", "key projects", "relevant projects", "portfolio",
                    "project portfolio", "project work", "project history", "project highlights", "major projects",
                    "significant projects", "project accomplishments", "project achievements", "project summary",
                    "project details", "project involvement", "project contributions"],
        "certifications": ["certifications", "certificates", "professional certifications", "credentials",
                          "professional credentials", "qualifications", "professional qualifications", "licenses",
                          "accreditations", "professional licenses", "professional accreditations", "certified",
                          "certification", "certificate", "credential", "license", "accreditation"],
        "awards": ["awards", "honors", "achievements", "recognitions", "accomplishments", "accolades",
                  "distinctions", "commendations", "acknowledgments", "prizes", "scholarships", "grants",
                  "fellowships", "recognition", "achievement", "honor", "award", "accomplishment"],
        "publications": ["publications", "papers", "research", "articles", "published works", "published papers",
                        "journal articles", "conference papers", "research papers", "academic publications",
                        "books", "book chapters", "technical papers", "white papers", "publication history"],
        "languages": ["languages", "language proficiency", "language skills", "foreign languages",
                     "language capabilities", "language competencies", "language knowledge", "language fluency",
                     "linguistic skills", "linguistic proficiency", "language abilities", "multilingual skills"],
        "interests": ["interests", "hobbies", "activities", "personal interests", "extracurricular activities",
                     "leisure activities", "pastimes", "recreational activities", "outside interests",
                     "personal activities", "personal pursuits", "personal hobbies", "personal passions"],
        "references": ["references", "professional references", "recommendations", "referees", "reference list",
                      "professional recommendations", "character references", "personal references",
                      "reference information", "references available", "references upon request"]
    }

    # Direct mapping for ALL CAPS headers (common in plain text CVs) - expanded for better matching
    direct_mapping = {
        "SUMMARY": "summary",
        "PROFILE": "summary",
        "OBJECTIVE": "summary",
        "ABOUT ME": "summary",
        "PERSONAL STATEMENT": "summary",
        "PERSONAL PROFILE": "summary",
        "CAREER SUMMARY": "summary",
        "EXECUTIVE SUMMARY": "summary",
        "OVERVIEW": "summary",
        "PROFESSIONAL OVERVIEW": "summary",
        "INTRODUCTION": "summary",

        "EXPERIENCE": "experience",
        "WORK EXPERIENCE": "experience",
        "PROFESSIONAL EXPERIENCE": "experience",
        "EMPLOYMENT HISTORY": "experience",
        "WORK HISTORY": "experience",
        "CAREER HISTORY": "experience",
        "PROFESSIONAL BACKGROUND": "experience",
        "EMPLOYMENT": "experience",
        "WORK": "experience",
        "CAREER": "experience",
        "JOBS": "experience",
        "POSITIONS HELD": "experience",

        "EDUCATION": "education",
        "ACADEMIC BACKGROUND": "education",
        "EDUCATIONAL BACKGROUND": "education",
        "QUALIFICATIONS": "education",
        "ACADEMIC QUALIFICATIONS": "education",
        "EDUCATIONAL QUALIFICATIONS": "education",
        "ACADEMIC HISTORY": "education",
        "SCHOOLING": "education",
        "ACADEMIC TRAINING": "education",
        "FORMAL EDUCATION": "education",

        "SKILLS": "skills",
        "TECHNICAL SKILLS": "skills",
        "CORE COMPETENCIES": "skills",
        "KEY SKILLS": "skills",
        "EXPERTISE": "skills",
        "PROFICIENCIES": "skills",
        "CAPABILITIES": "skills",
        "SKILL SET": "skills",
        "AREAS OF EXPERTISE": "skills",
        "COMPETENCIES": "skills",
        "STRENGTHS": "skills",
        "SPECIALTIES": "skills",
        "SPECIALIZATIONS": "skills",

        "PROJECTS": "projects",
        "PROJECT EXPERIENCE": "projects",
        "KEY PROJECTS": "projects",
        "RELEVANT PROJECTS": "projects",
        "PORTFOLIO": "projects",
        "PROJECT PORTFOLIO": "projects",
        "PROJECT WORK": "projects",
        "MAJOR PROJECTS": "projects",

        "CERTIFICATIONS": "certifications",
        "CERTIFICATES": "certifications",
        "PROFESSIONAL CERTIFICATIONS": "certifications",
        "CREDENTIALS": "certifications",
        "LICENSES": "certifications",
        "ACCREDITATIONS": "certifications",

        "AWARDS": "awards",
        "HONORS": "awards",
        "ACHIEVEMENTS": "awards",
        "RECOGNITIONS": "awards",
        "ACCOMPLISHMENTS": "awards",
        "ACCOLADES": "awards",

        "PUBLICATIONS": "publications",
        "PAPERS": "publications",
        "RESEARCH": "publications",
        "ARTICLES": "publications",
        "PUBLISHED WORKS": "publications",

        "LANGUAGES": "languages",
        "LANGUAGE PROFICIENCY": "languages",
        "LANGUAGE SKILLS": "languages",
        "FOREIGN LANGUAGES": "languages",

        "INTERESTS": "interests",
        "HOBBIES": "interests",
        "ACTIVITIES": "interests",
        "PERSONAL INTERESTS": "interests",
        "EXTRACURRICULAR ACTIVITIES": "interests",

        "REFERENCES": "references",
        "PROFESSIONAL REFERENCES": "references",
        "RECOMMENDATIONS": "references",
        "REFEREES": "references"
    }

    # Special approach for plain text CVs with ALL CAPS section headers
    # This is a common format in many CVs
    sections = {}
    lines = text.split('\n')

    # STRATEGY 1: Find potential section headers (all caps lines or lines with specific formatting)
    potential_headers = []
    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            continue

        # Check for ALL CAPS headers (very common in CVs)
        if line.isupper() and len(line) >= 3 and len(line) <= 35:
            # This looks like a section header
            if line in direct_mapping:
                section_name = direct_mapping[line]
                potential_headers.append((i, section_name, line))
                logging.info(f"Found ALL CAPS section header: '{line}' -> {section_name}")
            else:
                # Try to match it to a known section type
                matched_section = None
                for section_name, variations in section_headers.items():
                    if any(variation.lower() == line.lower() for variation in variations):
                        matched_section = section_name
                        break

                # If no exact match, try partial match
                if not matched_section:
                    for section_name, variations in section_headers.items():
                        if any(variation.lower() in line.lower() for variation in variations):
                            matched_section = section_name
                            break

                if matched_section:
                    potential_headers.append((i, matched_section, line))
                    logging.info(f"Found ALL CAPS section header (matched): '{line}' -> {matched_section}")

        # Check for Title Case headers with specific formatting (common in formatted CVs)
        elif (line.istitle() or line[0].isupper()) and len(line) >= 3 and len(line) <= 35:
            # This might be a section header
            matched_section = None

            # Check direct matches first
            for section_name, variations in section_headers.items():
                if any(variation.lower() == line.lower() for variation in variations):
                    matched_section = section_name
                    break

            # If no direct match, check for partial matches
            if not matched_section:
                for section_name, variations in section_headers.items():
                    if any(variation.lower() in line.lower() for variation in variations):
                        matched_section = section_name
                        break

            if matched_section:
                potential_headers.append((i, matched_section, line))
                logging.info(f"Found Title Case section header: '{line}' -> {matched_section}")

        # Check for headers with special formatting (e.g., followed by colon or with specific prefixes)
        elif ":" in line and len(line) <= 35:
            header_part = line.split(":")[0].strip()
            if header_part:
                matched_section = None
                for section_name, variations in section_headers.items():
                    if any(variation.lower() == header_part.lower() for variation in variations):
                        matched_section = section_name
                        break

                if not matched_section:
                    for section_name, variations in section_headers.items():
                        if any(variation.lower() in header_part.lower() for variation in variations):
                            matched_section = section_name
                            break

                if matched_section:
                    potential_headers.append((i, matched_section, line))
                    logging.info(f"Found colon-separated section header: '{line}' -> {matched_section}")

    # STRATEGY 2: Extract content between headers
    if len(potential_headers) >= 2:
        for i in range(len(potential_headers)):
            start_idx = potential_headers[i][0] + 1  # Start after the header
            end_idx = potential_headers[i+1][0] if i < len(potential_headers) - 1 else len(lines)
            section_name = potential_headers[i][1]

            # Extract content
            content = '\n'.join(lines[start_idx:end_idx]).strip()
            if content:
                sections[section_name] = content
                logging.info(f"Extracted {section_name} section with {len(content)} characters")

    # STRATEGY 3: If the above approach didn't work, try a more general approach with regex patterns
    if len(sections) < 2:
        # Reset sections
        sections = {}

        # Use regex patterns to find section headers and their content
        for section_name, variations in section_headers.items():
            for variation in variations:
                # Create pattern to match section headers with various formatting
                # This handles headers that are on their own line, followed by newline, or followed by colon
                pattern = re.compile(
                    rf'(?:^|\n)(?:\s*)(?:{re.escape(variation)}|{re.escape(variation.upper())}|{re.escape(variation.title())})(?:\s*:|\s*\n|\s*$)',
                    re.I
                )

                matches = list(pattern.finditer(text))
                if matches:
                    for match in matches:
                        start_pos = match.end()

                        # Find the end of this section (next section header or end of text)
                        end_pos = len(text)
                        for other_section, other_variations in section_headers.items():
                            for other_var in other_variations:
                                other_pattern = re.compile(
                                    rf'(?:^|\n)(?:\s*)(?:{re.escape(other_var)}|{re.escape(other_var.upper())}|{re.escape(other_var.title())})(?:\s*:|\s*\n|\s*$)',
                                    re.I
                                )
                                other_matches = list(other_pattern.finditer(text[start_pos:]))
                                if other_matches:
                                    potential_end = start_pos + other_matches[0].start()
                                    if potential_end < end_pos:
                                        end_pos = potential_end

                        # Extract the section content
                        section_content = text[start_pos:end_pos].strip()
                        if section_content and len(section_content) > 15:  # Avoid tiny sections
                            sections[section_name] = section_content
                            logging.info(f"Extracted {section_name} section with regex: {len(section_content)} characters")

                        # Only use the first match for each variation to avoid duplicates
                        break

                    # If we found content for this section, move to the next section
                    if section_name in sections:
                        break

    # STRATEGY 4: If we still don't have enough sections, try a line-by-line approach
    if len(sections) < 2:
        # Reset sections
        sections = {}
        current_section = None
        section_content = []

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Check if this line is a section header
            is_header = False
            for section_name, header_variations in section_headers.items():
                # More flexible pattern matching
                if any(header.lower() in line.lower() and (line.isupper() or line.istitle() or re.search(r'^[A-Z]', line))
                       for header in header_variations):
                    # If we were building a section, save it
                    if current_section and section_content:
                        sections[current_section] = '\n'.join(section_content)

                    # Start a new section
                    current_section = section_name
                    section_content = []
                    is_header = True
                    logging.info(f"Found section header (line-by-line): '{line}' -> {section_name}")
                    break

            # If not a header, add to current section content
            if not is_header and current_section:
                section_content.append(line)

        # Save the last section
        if current_section and section_content:
            sections[current_section] = '\n'.join(section_content)
            logging.info(f"Saved final section: {current_section} with {len(section_content)} lines")

    # STRATEGY 5: If we still don't have key sections, try to infer them from content
    if "skills" not in sections:
        # Look for skill-like content (bullet points, comma-separated lists)
        skill_patterns = [
            r'(?:^|\n)(?:\s*[-•*]\s*|\d+\.\s*)([A-Za-z][A-Za-z\s]+(?:development|programming|design|management|engineering|analysis|communication|languages|tools|frameworks|databases|platforms|systems|methodologies))',
            r'(?:^|\n)(?:skills|expertise|competencies|proficiencies)(?:\s*:|\s*include|\s*are)?\s*([A-Za-z][A-Za-z\s,]+)'
        ]

        for pattern in skill_patterns:
            skill_matches = re.findall(pattern, text, re.I)
            if skill_matches:
                sections["skills"] = "\n".join(skill_matches)
                logging.info(f"Inferred skills section with {len(skill_matches)} matches")
                break

    if "education" not in sections:
        # Look for education-like content (degrees, universities)
        edu_patterns = [
            r'(?:^|\n)(?:\s*[-•*]\s*|\d+\.\s*)?([A-Za-z][A-Za-z\s]+(?:University|College|Institute|School))',
            r'(?:^|\n)(?:\s*[-•*]\s*|\d+\.\s*)?(?:Bachelor|Master|PhD|Doctorate|BSc|MSc|BA|MA|MBA|Degree|Diploma)'
        ]

        for pattern in edu_patterns:
            edu_matches = re.findall(pattern, text, re.I)
            if edu_matches:
                sections["education"] = "\n".join(edu_matches)
                logging.info(f"Inferred education section with {len(edu_matches)} matches")
                break

    # Log the sections we found
    logging.info(f"Found {len(sections)} sections: {list(sections.keys())}")

    return sections

def get_section_content(sections: Dict[str, str], section_names: List[str]) -> str:
    """
    Get content from a section, trying multiple possible section names.

    This improved version:
    1. Tries multiple section name variations
    2. Handles case-insensitive matching
    3. Cleans sensitive information from the content
    4. Preserves formatting for better readability
    5. Handles empty or malformed sections
    6. Combines content from multiple matching sections if needed

    Args:
        sections: Dictionary of section name -> content
        section_names: List of possible section names to try

    Returns:
        Cleaned section content or empty string if not found
    """
    combined_content = []

    # Try each section name
    for name in section_names:
        # Try exact match first
        if name in sections and sections[name]:
            combined_content.append(sections[name])
        else:
            # Try case-insensitive match
            for section_key in sections:
                if section_key.lower() == name.lower() and sections[section_key]:
                    combined_content.append(sections[section_key])
                    break

    # If we found content, clean and return it
    if combined_content:
        # Join multiple sections if we found more than one
        content = "\n\n".join(combined_content)

        # Clean sensitive information
        content = clean_sensitive_information(content)

        # Preserve paragraph breaks for better readability
        paragraphs = [p.strip() for p in content.split('\n') if p.strip()]
        content = '\n'.join(paragraphs)

        return content

    # If no content found, return empty string
    return ""

def extract_skills(text: str, skills_section: str, doc) -> List[str]:
    """
    Enhanced skill extraction using multiple strategies:
    1. Look for skills in the dedicated skills section
    2. Use NER and pattern matching throughout the document
    3. Match against comprehensive skill dictionary
    4. Detect skill proficiency levels
    5. Categorize skills by type
    6. Provide context for each skill
    7. Handle bullet points, comma-separated lists, and other common formats
    8. Detect domain-specific skills not in our dictionary
    9. Analyze skill usage context to determine proficiency

    Returns a list of skill names, sorted by confidence and relevance.
    """
    found_skills = {}  # Dictionary to store skill info: {skill_name: {category, proficiency, context}}

    # Proficiency level indicators and their scores (0-1 scale)
    proficiency_indicators = {
        'expert': 1.0,
        'advanced': 0.9,
        'proficient': 0.8,
        'experienced': 0.8,
        'skilled': 0.7,
        'intermediate': 0.6,
        'familiar': 0.5,
        'basic': 0.4,
        'beginner': 0.3,
        'novice': 0.2,
        'learning': 0.1,
        'exposure': 0.1
    }

    # STRATEGY 1: Extract from skills section using comprehensive skill list
    if skills_section:
        # First, look for structured skill lists (bullet points, comma-separated)
        bullet_pattern = r'(?:^|\n)(?:\s*[-•*]\s*|\d+\.\s*)([^\n]+)'
        bullet_points = re.findall(bullet_pattern, skills_section)

        # Process bullet points in skills section (highest confidence)
        for point in bullet_points:
            # Check for skills in our dictionary
            for skill in ALL_SKILLS:
                if re.search(rf'\b{re.escape(skill)}\b', point, re.I):
                    # Determine proficiency from the bullet point
                    proficiency = 0.7  # Default medium-high proficiency for skills section

                    # Check for proficiency indicators
                    if re.search(r'\b(?:expert|advanced|proficient|strong|extensive|mastery)\b', point, re.I):
                        proficiency = 0.9
                    elif re.search(r'\b(?:intermediate|working knowledge|good|solid|competent)\b', point, re.I):
                        proficiency = 0.7
                    elif re.search(r'\b(?:basic|beginner|familiar|exposure|learning|novice)\b', point, re.I):
                        proficiency = 0.4

                    skill_info = {
                        'category': SKILL_CATEGORIES.get(skill, "Other"),
                        'proficiency': proficiency,
                        'context': 'Listed in skills section: ' + point.strip(),
                        'confidence': 0.95  # Highest confidence for skills explicitly listed
                    }
                    found_skills[skill] = skill_info

            # Also look for potential skills not in our dictionary
            # These might be domain-specific or newer technologies
            potential_skill = point.strip()

            # Clean up the potential skill
            potential_skill = re.sub(r'^[-•*\d.]+\s*', '', potential_skill)  # Remove bullet markers
            potential_skill = re.sub(r'\s*[-:]\s*.*$', '', potential_skill)  # Remove descriptions after dash/colon

            # Check if it looks like a valid skill (not too long, not a sentence)
            if (len(potential_skill) >= 2 and len(potential_skill) <= 50 and
                len(potential_skill.split()) <= 5 and  # Not too many words
                not re.search(r'\b(?:and|or|the|with|using|for|to|in|on|by|at|from|as)\b', potential_skill, re.I)):  # Not a phrase

                # Skip if it's already in our found skills
                if potential_skill.lower() not in [s.lower() for s in found_skills.keys()]:
                    skill_info = {
                        'category': "Domain-Specific Skill",
                        'proficiency': 0.7,
                        'context': 'Listed in skills section',
                        'confidence': 0.8  # High confidence but not as high as known skills
                    }
                    found_skills[potential_skill] = skill_info

        # Look for comma-separated skill lists
        comma_lists = re.findall(r'(?:^|\n)([^-•*\n][^:]+?(?::|include|are|in)?\s*([A-Za-z][A-Za-z\s,/&+#]+))', skills_section)
        for full_match, skill_list in comma_lists:
            if ',' in skill_list:  # Only process if it looks like a comma-separated list
                skills_items = [s.strip() for s in skill_list.split(',')]
                for item in skills_items:
                    if len(item) >= 2 and len(item) <= 50:  # Reasonable length for a skill
                        # Check against our dictionary first
                        matched = False
                        for skill in ALL_SKILLS:
                            if re.search(rf'\b{re.escape(skill)}\b', item, re.I):
                                skill_info = {
                                    'category': SKILL_CATEGORIES.get(skill, "Other"),
                                    'proficiency': 0.7,
                                    'context': 'Listed in skills section',
                                    'confidence': 0.9
                                }
                                found_skills[skill] = skill_info
                                matched = True

                        # If not in dictionary, add as potential skill
                        if not matched and item.lower() not in [s.lower() for s in found_skills.keys()]:
                            skill_info = {
                                'category': "Domain-Specific Skill",
                                'proficiency': 0.7,
                                'context': 'Listed in skills section',
                                'confidence': 0.75
                            }
                            found_skills[item] = skill_info

        # Also do a general search for known skills in the skills section
        for skill in ALL_SKILLS:
            # Skip if already found
            if skill in found_skills:
                continue

            # Use word boundary to avoid partial matches
            if re.search(rf'\b{re.escape(skill)}\b', skills_section, re.I):
                skill_info = {
                    'category': SKILL_CATEGORIES.get(skill, "Other"),
                    'proficiency': 0.7,  # Default medium-high proficiency for skills section
                    'context': 'Listed in skills section',
                    'confidence': 0.85
                }
                found_skills[skill] = skill_info

    # STRATEGY 2: Look for skills in bullet points throughout the document
    bullet_pattern = r'(?:^|\n)(?:\s*[-•*]\s*|\d+\.\s*)([^\n]+)'
    bullet_points = re.findall(bullet_pattern, text)

    for point in bullet_points:
        for skill in ALL_SKILLS:
            if re.search(rf'\b{re.escape(skill)}\b', point, re.I):
                # If already found in skills section, skip
                if skill in found_skills:
                    continue

                # Determine context and proficiency from the bullet point
                proficiency = 0.6  # Default medium proficiency for bullet points

                # Check for proficiency indicators
                if re.search(r'\b(?:expert|advanced|proficient|strong|extensive)\b', point, re.I):
                    proficiency = 0.9
                elif re.search(r'\b(?:intermediate|working knowledge|good|solid)\b', point, re.I):
                    proficiency = 0.7
                elif re.search(r'\b(?:basic|beginner|familiar|exposure|learning)\b', point, re.I):
                    proficiency = 0.4

                # Check for skill usage context
                context_score = 0.7  # Default
                if re.search(r'\b(?:developed|implemented|created|built|designed|architected|managed|led)\b', point, re.I):
                    context_score = 0.85  # Higher confidence for active usage

                skill_info = {
                    'category': SKILL_CATEGORIES.get(skill, "Other"),
                    'proficiency': proficiency,
                    'context': point[:100] + ('...' if len(point) > 100 else ''),
                    'confidence': context_score
                }
                found_skills[skill] = skill_info

    # STRATEGY 3: Look for skill patterns with proficiency indicators
    proficiency_patterns = [
        # Pattern: "Proficient in X" or "Advanced X skills"
        r'(?P<level>expert|advanced|proficient|experienced|skilled|intermediate|familiar|basic|beginner|novice)\s+(?:in|with|using)?\s+(?P<skill>[A-Za-z0-9+#\s]+)',
        r'(?P<level>expert|advanced|proficient|experienced|skilled|intermediate|familiar|basic|beginner|novice)\s+(?P<skill>[A-Za-z0-9+#\s]+)\s+(?:skills|knowledge|proficiency|expertise)',
        r'(?P<skill>[A-Za-z0-9+#\s]+)\s+(?P<level>expert|advanced|proficient|experienced|skilled|intermediate|familiar|basic|beginner|novice)',
    ]

    for pattern in proficiency_patterns:
        for match in re.finditer(pattern, text, re.I):
            level = match.group('level').lower()
            potential_skill = match.group('skill').strip().lower()
            proficiency_score = proficiency_indicators.get(level, 0.5)

            # Get the context (sentence containing the skill)
            context_start = max(0, match.start() - 50)
            context_end = min(len(text), match.end() + 50)
            context = text[context_start:context_end].strip()

            # Check if it's in our skill list or a substring of a skill
            for skill in ALL_SKILLS:
                if potential_skill == skill.lower() or potential_skill in skill.lower():
                    # If we already found this skill, update proficiency if higher
                    if skill in found_skills:
                        if proficiency_score > found_skills[skill]['proficiency']:
                            found_skills[skill]['proficiency'] = proficiency_score
                            found_skills[skill]['context'] = context
                    else:
                        skill_info = {
                            'category': SKILL_CATEGORIES.get(skill, "Other"),
                            'proficiency': proficiency_score,
                            'context': context,
                            'confidence': 0.8  # Higher confidence due to explicit proficiency
                        }
                        found_skills[skill] = skill_info

    # STRATEGY 4: Look for general skill mentions
    general_skill_patterns = [
        r'(?:skills|expertise|proficiency|competency|knowledge)\s+(?:in|with|using)?\s+([A-Za-z0-9+#\s,]+)',
        r'([A-Za-z0-9+#\s]+)\s+(?:skills|expertise|proficiency|experience)'
    ]

    for pattern in general_skill_patterns:
        matches = re.findall(pattern, text, re.I)
        for match in matches:
            # Clean up the match and split by commas for lists
            potential_skills = [s.strip().lower() for s in re.split(r'[,;]', match) if s.strip()]
            for potential_skill in potential_skills:
                # Check if it's in our skill list or a substring of a skill
                for skill in ALL_SKILLS:
                    if potential_skill == skill.lower() or potential_skill in skill.lower():
                        if skill not in found_skills:
                            skill_info = {
                                'category': SKILL_CATEGORIES.get(skill, "Other"),
                                'proficiency': 0.6,  # Default medium proficiency
                                'context': 'Mentioned as a skill',
                                'confidence': 0.7
                            }
                            found_skills[skill] = skill_info

    # STRATEGY 5: Look for skills in experience section
    experience_section = get_section_content(extract_sections(text), ["experience", "work experience", "professional experience"])
    if experience_section:
        for skill in ALL_SKILLS:
            if re.search(rf'\b{re.escape(skill)}\b', experience_section, re.I):
                if skill not in found_skills:
                    skill_info = {
                        'category': SKILL_CATEGORIES.get(skill, "Other"),
                        'proficiency': 0.7,  # Higher proficiency for skills in experience
                        'context': 'Used in professional experience',
                        'confidence': 0.8
                    }
                    found_skills[skill] = skill_info
                else:
                    # Boost confidence for skills found in both skills and experience sections
                    found_skills[skill]['confidence'] = max(found_skills[skill]['confidence'], 0.9)
                    if 'Listed in skills section' in found_skills[skill]['context']:
                        found_skills[skill]['context'] = 'Listed in skills section and used in professional experience'

    # STRATEGY 6: Look for programming languages, frameworks, and technologies in code-like contexts
    code_patterns = [
        r'\b(?:using|import|require|include|from)\s+([A-Za-z][A-Za-z0-9_.-]+)',
        r'\b(?:programmed|coded|developed|implemented)\s+(?:in|with|using)?\s+([A-Za-z][A-Za-z0-9_.-]+)',
        r'<([A-Za-z][A-Za-z0-9_.-]+)>',  # HTML/XML tags
        r'([A-Za-z][A-Za-z0-9_.-]+)(?:\.[A-Za-z][A-Za-z0-9_.-]+)+',  # Object.method patterns
    ]

    for pattern in code_patterns:
        matches = re.findall(pattern, text)
        for match in matches:
            potential_tech = match.lower()
            # Check if it's a known technology
            for skill in PROGRAMMING_LANGUAGES + WEB_TECHNOLOGIES + DATA_SCIENCE + CLOUD_DEVOPS + DATABASE_TECHNOLOGIES:
                if potential_tech == skill.lower() or potential_tech.startswith(skill.lower() + '.'):
                    if skill not in found_skills:
                        skill_info = {
                            'category': SKILL_CATEGORIES.get(skill, "Technology"),
                            'proficiency': 0.7,
                            'context': 'Found in code-like context',
                            'confidence': 0.8
                        }
                        found_skills[skill] = skill_info

    # STRATEGY 7: Use NER to identify technical terms not in our dictionary
    try:
        # Only process first 50K chars to avoid memory issues
        doc_sample = doc(text[:50000]) if len(text) > 50000 else doc

        for ent in doc_sample.ents:
            # Look for entities that might be skills (ORG can be technologies, PRODUCT can be software)
            if ent.label_ in ["ORG", "PRODUCT", "GPE"]:
                skill_name = ent.text.lower()

                # Skip common non-skill entities and very short names
                if (len(skill_name) < 3 or
                    skill_name in ["company", "organization", "university", "college", "school",
                                  "corporation", "inc", "llc", "ltd", "gmbh"]):
                    continue

                # Skip if already in our dictionary
                if skill_name in found_skills:
                    continue

                # Check if it might be a technical skill
                if (re.search(r'\b(?:software|framework|library|platform|tool|language|database|system|technology|api)\b',
                             ent.sent.text, re.I) or
                    any(tech in ent.sent.text.lower() for tech in ["developed", "implemented", "programmed",
                                                                  "designed", "architected", "built", "created"])):

                    skill_info = {
                        'category': "Detected Technology",
                        'proficiency': 0.6,
                        'context': ent.sent.text[:100] + ('...' if len(ent.sent.text) > 100 else ''),
                        'confidence': 0.5
                    }
                    found_skills[skill_name] = skill_info
    except Exception as e:
        logging.warning(f"Error in NER skill extraction: {e}")

    # STRATEGY 8: Look for soft skills in context
    soft_skill_patterns = [
        r'\b(?:strong|excellent|good|proven|demonstrated)\s+([A-Za-z][A-Za-z\s]+?)\s+(?:skills|abilities|capabilities)',
        r'\b(?:skilled|proficient|experienced)\s+(?:in|with)?\s+([A-Za-z][A-Za-z\s]+?)\s+(?:and|,|\.|$)',
    ]

    for pattern in soft_skill_patterns:
        matches = re.findall(pattern, text, re.I)
        for match in matches:
            potential_skill = match.lower()
            # Check if it's a known soft skill
            for skill in SOFT_SKILLS + BUSINESS_SKILLS:
                if potential_skill == skill.lower() or skill.lower() in potential_skill:
                    if skill not in found_skills:
                        skill_info = {
                            'category': SKILL_CATEGORIES.get(skill, "Soft Skill"),
                            'proficiency': 0.7,
                            'context': 'Found in soft skill context',
                            'confidence': 0.7
                        }
                        found_skills[skill] = skill_info

    # Convert to list of skill names, sorted by confidence and proficiency
    skills_list = []
    for skill_name, info in sorted(found_skills.items(), key=lambda x: (x[1]['confidence'], x[1]['proficiency']), reverse=True):
        # Format skill name appropriately
        if len(skill_name) <= 3:  # Acronyms like CSS, AWS
            formatted_skill = skill_name.upper()
        elif '.' in skill_name and not ' ' in skill_name:  # Technology with dots like Node.js
            parts = skill_name.split('.')
            formatted_skill = '.'.join(p.capitalize() for p in parts)
        else:
            # Capitalize each word for multi-word skills
            formatted_skill = ' '.join(word.capitalize() if len(word) > 3 else word
                                     for word in skill_name.split())

        skills_list.append({
            'name': formatted_skill,
            'category': info['category'],
            'proficiency': round(info['proficiency'], 2),
            'context': info['context'],
            'confidence': round(info['confidence'], 2)
        })

    # Limit to top 30 skills to avoid overwhelming the user
    if len(skills_list) > 30:
        skills_list = skills_list[:30]

    # For backward compatibility, return a simple list of skill names
    skill_names = [skill['name'] for skill in skills_list]

    return skill_names

def extract_education(text: str, doc) -> List[Dict[str, str]]:
    """
    Extract structured education information from the education section.
    Returns a list of education entries, each with degree, institution, dates, location, and GPA.

    This improved version uses multiple strategies to accurately identify education information:
    1. Looks for degree keywords and patterns
    2. Identifies educational institutions
    3. Extracts dates and locations
    4. Detects GPA and academic achievements
    5. Uses NER to identify organizations and locations
    6. Handles various formatting styles (bullet points, paragraphs)
    """
    if not text:
        return []

    logging.info(f"Extracting education from text: {text[:100]}...")

    # Define comprehensive patterns for education extraction
    degree_patterns = [
        # Bachelor's degree patterns
        r'(?:bachelor|bachelors|baccalaureate|b\.?s\.?|b\.?a\.?|b\.?eng\.?|b\.?tech\.?|b\.?sc\.?|b\.?b\.?a\.?|undergraduate)\s+(?:of|in|degree)?\s+([A-Za-z\s,&]+)',
        r'([A-Za-z\s,&]+)\s+(?:bachelor|bachelors|baccalaureate)',

        # Master's degree patterns
        r'(?:master|masters|m\.?s\.?|m\.?a\.?|m\.?eng\.?|m\.?tech\.?|m\.?sc\.?|m\.?b\.?a\.?|graduate)\s+(?:of|in|degree)?\s+([A-Za-z\s,&]+)',
        r'([A-Za-z\s,&]+)\s+(?:master|masters)',

        # Doctorate patterns
        r'(?:phd|ph\.?d\.?|doctorate|doctor|d\.?sc\.?|d\.?phil\.?)\s+(?:of|in|degree)?\s+([A-Za-z\s,&]+)',
        r'([A-Za-z\s,&]+)\s+(?:phd|ph\.?d\.?|doctorate)',

        # Associate and other degrees
        r'(?:associate|diploma|certificate|a\.?a\.?|a\.?s\.?)\s+(?:of|in|degree)?\s+([A-Za-z\s,&]+)',
        r'([A-Za-z\s,&]+)\s+(?:associate|diploma|certificate)',

        # General degree patterns
        r'(?:degree|qualification)\s+(?:in|of)\s+([A-Za-z\s,&]+)',
        r'([A-Za-z\s,&]+)\s+(?:degree|qualification)'
    ]

    institution_patterns = [
        # University patterns
        r'(?:university|college|institute|school|academy|polytechnic)\s+(?:of)?\s+([A-Za-z\s,&]+)',
        r'([A-Za-z\s,&]+)\s+(?:university|college|institute|school|academy|polytechnic)',

        # Common institution naming patterns
        r'([A-Z][A-Za-z\s,&]+)(?:university|college|institute|school|academy|polytechnic)',
        r'([A-Z][A-Za-z\s,&]+)(?:\s+in\s+[A-Za-z\s,&]+)',

        # Institutions with locations
        r'(?:university|college|institute|school|academy|polytechnic)\s+(?:at|in|of)\s+([A-Za-z\s,&]+)',

        # Abbreviated institutions
        r'\b([A-Z][A-Z]+)\b'  # Potential institution abbreviations
    ]

    date_patterns = [
        # Month-Year to Month-Year format
        r'(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s+\d{4}\s*(?:-|–|to)\s*(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s+\d{4}',

        # Month-Year to Present format
        r'(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s+\d{4}\s*(?:-|–|to)\s*(?:present|current|now|ongoing)',

        # Year to Year format
        r'\b\d{4}\s*(?:-|–|to)\s*\d{4}\b',

        # Year to Present format
        r'\b\d{4}\s*(?:-|–|to)\s*(?:present|current|now|ongoing)\b',

        # Single Month-Year format (graduation date)
        r'(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s+\d{4}',

        # Single Year format (graduation date)
        r'(?:graduated|completed|finished|earned|received|conferred|awarded)(?:\s+in)?\s+(\d{4})',
        r'(?:class|batch|year)(?:\s+of)?\s+(\d{4})',
        r'\b(?:19|20)\d{2}\b'  # Years from 1900-2099
    ]

    location_patterns = [
        # Location with context
        r'(?:located|located in|in|at)\s+([A-Za-z\s,]+)',

        # City, State/Country format
        r'([A-Za-z\s]+),\s*(?:[A-Z]{2}|[A-Za-z]+\s+[A-Za-z]+)',

        # Common location patterns in education
        r'(?:campus|branch|center)\s+(?:in|at)\s+([A-Za-z\s,]+)'
    ]

    gpa_patterns = [
        # GPA with scale
        r'(?:gpa|grade point average|g\.p\.a\.)(?:\s+of)?\s+(\d+\.\d+)(?:/\d+\.\d+)?',
        r'(\d+\.\d+)(?:/\d+\.\d+)?\s+(?:gpa|grade point average|g\.p\.a\.)',

        # Percentage grades
        r'(?:average|grade|score|percentage)(?:\s+of)?\s+(\d+\.?\d*)\s*%',

        # First class, second class honors
        r'(first|second|third)(?:\s+class)?\s+(?:honors|honours|distinction)',

        # Cum laude
        r'(summa|magna|cum)\s+laude',

        # General academic standing
        r'(?:with)?\s+(honors|honours|distinction|merit)'
    ]

    # Expanded degree keywords for better matching
    degree_keywords = [
        "bachelor", "bachelors", "baccalaureate", "bs", "ba", "bsc", "beng", "btech", "bba", "undergraduate",
        "master", "masters", "ms", "ma", "msc", "meng", "mtech", "mba", "graduate",
        "phd", "ph.d", "doctorate", "doctor", "dsc", "dphil",
        "associate", "aa", "as", "diploma", "certificate",
        "degree", "qualification"
    ]

    # Expanded institution keywords
    institution_keywords = [
        "university", "college", "institute", "school", "academy", "polytechnic",
        "campus", "department", "faculty"
    ]

    # First, try to identify distinct education entries by looking for degree keywords and institutions
    education_entries = []
    lines = text.split('\n')

    # STRATEGY 1: Group lines into entries based on degree keywords and institutions
    current_entry = []
    entries = []

    for line in lines:
        line = line.strip()
        if not line:
            if current_entry:
                entries.append('\n'.join(current_entry))
                current_entry = []
            continue

        # Check if this line starts a new entry (contains a degree keyword or institution)
        is_new_entry = False

        # Check for degree keywords
        if any(keyword in line.lower() for keyword in degree_keywords):
            is_new_entry = True

        # Check for institution keywords
        if not is_new_entry and any(keyword in line.lower() for keyword in institution_keywords):
            is_new_entry = True

        # Check for date patterns (often indicate the start of a new entry)
        if not is_new_entry:
            for pattern in date_patterns:
                if re.search(pattern, line, re.I):
                    is_new_entry = True
                    break

        if is_new_entry:
            if current_entry:
                entries.append('\n'.join(current_entry))
                current_entry = []
            current_entry.append(line)
        elif current_entry:
            current_entry.append(line)

    # Add the last entry if there is one
    if current_entry:
        entries.append('\n'.join(current_entry))

    # STRATEGY 2: If we couldn't identify distinct entries, try splitting by formatting cues
    if not entries:
        # Try splitting by bullet points, blank lines, or numbering
        entries = re.split(r'\n\s*\n|\n\s*•|\n\s*-|\n\s*\*|\n\s*\d+\.', text)

        # If still no entries, treat the whole text as one entry
        if not entries or (len(entries) == 1 and not entries[0].strip()):
            entries = [text]

    # STRATEGY 3: Process each entry to extract structured information
    for entry in entries:
        if not entry.strip():
            continue

        education_entry = {
            'degree': '',
            'institution': '',
            'dates': '',
            'location': '',
            'gpa': '',
            'description': entry.strip()
        }

        # Extract degree
        for pattern in degree_patterns:
            degree_matches = re.findall(pattern, entry, re.I)
            if degree_matches:
                # Get the full match context to include the degree type
                match_context = re.search(pattern, entry, re.I)
                if match_context:
                    # Try to get the full degree (type + field)
                    full_match = match_context.group(0).strip()
                    if len(full_match) <= 100:  # Reasonable length for a degree
                        education_entry['degree'] = full_match
                    else:
                        # If too long, just use the field of study
                        education_entry['degree'] = degree_matches[0].strip()
                else:
                    education_entry['degree'] = degree_matches[0].strip()
                break

        # If no degree found with patterns, look for degree keywords in context
        if not education_entry['degree']:
            for keyword in degree_keywords:
                if keyword in entry.lower():
                    # Find the sentence containing the keyword
                    sentences = re.split(r'[.!?]\s+', entry)
                    for sentence in sentences:
                        if keyword in sentence.lower():
                            # Extract a reasonable portion of the sentence
                            words = sentence.split()
                            keyword_index = -1
                            for i, word in enumerate(words):
                                if keyword.lower() in word.lower():
                                    keyword_index = i
                                    break

                            if keyword_index >= 0:
                                # Get a window around the keyword (up to 7 words)
                                start_idx = max(0, keyword_index - 3)
                                end_idx = min(len(words), keyword_index + 4)
                                degree_phrase = ' '.join(words[start_idx:end_idx])
                                education_entry['degree'] = degree_phrase.strip()
                                break
                    if education_entry['degree']:
                        break

        # Extract institution
        for pattern in institution_patterns:
            institution_matches = re.findall(pattern, entry, re.I)
            if institution_matches:
                # Get the full match context to include the institution type
                match_context = re.search(pattern, entry, re.I)
                if match_context:
                    # Try to get the full institution name
                    full_match = match_context.group(0).strip()
                    if len(full_match) <= 100:  # Reasonable length for an institution
                        education_entry['institution'] = full_match
                    else:
                        # If too long, just use the institution name
                        education_entry['institution'] = institution_matches[0].strip()
                else:
                    education_entry['institution'] = institution_matches[0].strip()
                break

        # Extract dates
        for pattern in date_patterns:
            date_matches = re.findall(pattern, entry, re.I)
            if date_matches:
                if isinstance(date_matches[0], tuple):
                    # Some patterns return tuples, get the first element
                    education_entry['dates'] = date_matches[0][0].strip()
                else:
                    education_entry['dates'] = date_matches[0].strip()
                break

        # Extract location
        for pattern in location_patterns:
            location_matches = re.findall(pattern, entry, re.I)
            if location_matches:
                education_entry['location'] = location_matches[0].strip()
                break

        # Extract GPA and academic achievements
        for pattern in gpa_patterns:
            gpa_matches = re.findall(pattern, entry, re.I)
            if gpa_matches:
                if isinstance(gpa_matches[0], tuple):
                    education_entry['gpa'] = gpa_matches[0][0].strip()
                else:
                    education_entry['gpa'] = gpa_matches[0].strip()
                break

        # STRATEGY 4: Use NER to identify organizations and locations if not found with patterns
        if not education_entry['institution'] or not education_entry['location']:
            try:
                entry_doc = doc(entry)
                for ent in entry_doc.ents:
                    if ent.label_ == "ORG" and not education_entry['institution']:
                        # Check if it looks like an educational institution
                        if any(keyword in ent.text.lower() for keyword in institution_keywords):
                            education_entry['institution'] = ent.text
                        # Or if it's capitalized and not too long
                        elif ent.text[0].isupper() and len(ent.text.split()) <= 5:
                            education_entry['institution'] = ent.text
                    elif ent.label_ == "GPE" and not education_entry['location']:
                        education_entry['location'] = ent.text
            except Exception as e:
                logging.warning(f"Error using NER for education extraction: {e}")

        # STRATEGY 5: Clean up and validate the extracted information

        # Clean up degree (remove extra whitespace, normalize common abbreviations)
        if education_entry['degree']:
            # Normalize common degree abbreviations
            degree_text = education_entry['degree']
            degree_text = re.sub(r'\bB\.?S\.?\b', 'Bachelor of Science', degree_text, flags=re.I)
            degree_text = re.sub(r'\bB\.?A\.?\b', 'Bachelor of Arts', degree_text, flags=re.I)
            degree_text = re.sub(r'\bM\.?S\.?\b', 'Master of Science', degree_text, flags=re.I)
            degree_text = re.sub(r'\bM\.?A\.?\b', 'Master of Arts', degree_text, flags=re.I)
            degree_text = re.sub(r'\bM\.?B\.?A\.?\b', 'Master of Business Administration', degree_text, flags=re.I)
            degree_text = re.sub(r'\bPh\.?D\.?\b', 'Doctor of Philosophy', degree_text, flags=re.I)

            # Clean up whitespace
            degree_text = re.sub(r'\s+', ' ', degree_text).strip()
            education_entry['degree'] = degree_text

        # Clean up institution name
        if education_entry['institution']:
            # Remove any trailing punctuation
            institution_text = re.sub(r'[,;.]\s*$', '', education_entry['institution']).strip()
            education_entry['institution'] = institution_text

        # Only add entries that have at least a degree or institution
        if education_entry['degree'] or education_entry['institution']:
            education_entries.append(education_entry)

    # STRATEGY 6: If we still don't have any entries but have text, try a more aggressive approach
    if not education_entries and text.strip():
        # Look for any mentions of educational institutions
        for pattern in institution_patterns:
            institution_matches = re.findall(pattern, text, re.I)
            if institution_matches:
                for match in institution_matches:
                    education_entry = {
                        'degree': '',
                        'institution': match.strip(),
                        'dates': '',
                        'location': '',
                        'gpa': '',
                        'description': text.strip()
                    }

                    # Try to find a degree near this institution
                    context_start = max(0, text.find(match) - 100)
                    context_end = min(len(text), text.find(match) + len(match) + 100)
                    context = text[context_start:context_end]

                    for keyword in degree_keywords:
                        if keyword in context.lower():
                            # Find the sentence containing the keyword
                            sentences = re.split(r'[.!?]\s+', context)
                            for sentence in sentences:
                                if keyword in sentence.lower():
                                    education_entry['degree'] = sentence.strip()
                                    break
                            break

                    education_entries.append(education_entry)
                    break
            if education_entries:
                break

    logging.info(f"Extracted {len(education_entries)} education entries")
    return education_entries

def extract_experience(text: str, doc) -> List[Dict[str, str]]:
    """
    Extract structured work experience information.
    Returns a list of experience entries, each with title, company, dates, location, description, and achievements.

    This improved version uses multiple strategies to accurately identify work experience:
    1. Looks for job titles and company names
    2. Identifies employment dates and locations
    3. Extracts achievements and responsibilities
    4. Uses NER to identify organizations and locations
    5. Handles various formatting styles (bullet points, paragraphs)
    6. Detects job transitions and multiple positions
    """
    if not text:
        return []

    logging.info(f"Extracting experience from text: {text[:100]}...")

    # Define comprehensive patterns for experience extraction
    title_patterns = [
        # Job title followed by company indicator
        r'(?:^|\n)(?:\s*[-•*]\s*)?([A-Z][A-Za-z\s&,\-]+)(?:\n|,|\s+at\s+|\s+for\s+|\s+with\s+)',

        # Job title with explicit markers
        r'(?:^|\n)(?:as|position|title|role)(?:\s+of)?(?:\s+a)?(?:\s+the)?\s+([A-Za-z\s&,\-]+)',
        r'(?:job|position|title|role)(?:\s*:|\s+is|\s+was)?\s+([A-Za-z\s&,\-]+)',

        # Job title in context
        r'(?:hired|employed|worked|joined)(?:\s+as)?\s+(?:a|an)?\s+([A-Za-z\s&,\-]+)',
        r'(?:promoted|advanced|moved)\s+to\s+(?:a|an)?\s+([A-Za-z\s&,\-]+)',

        # Job title with common suffixes
        r'([A-Za-z\s&,\-]+)(?:\s+position|\s+role|\s+title)'
    ]

    company_patterns = [
        # Company with prepositions
        r'(?:at|with|for)\s+([A-Z][A-Za-z\s&,\-\.]+)',

        # Company with explicit markers
        r'(?:company|employer|organization|firm|business|corporation)(?:\s+name)?(?:\s*:|\s+is|\s+was)?\s+([A-Z][A-Za-z\s&,\-\.]+)',

        # Company in context
        r'(?:joined|worked|employed)(?:\s+(?:at|with|for|by))?\s+([A-Z][A-Za-z\s&,\-\.]+)',

        # Company followed by punctuation or line break
        r'([A-Z][A-Za-z\s&,\-\.]+)(?:,|\.|\n)',

        # Company with location
        r'([A-Z][A-Za-z\s&,\-\.]+)(?:\s+in\s+[A-Za-z\s,]+)',

        # Company with LLC, Inc, etc.
        r'([A-Z][A-Za-z\s&,\-\.]+)(?:\s+(?:LLC|Inc|Ltd|Limited|Corporation|Corp|GmbH))'
    ]

    date_patterns = [
        # Month-Year to Month-Year format
        r'(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s+\d{4}\s*(?:-|–|to)\s*(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s+\d{4}',

        # Month-Year to Present format
        r'(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s+\d{4}\s*(?:-|–|to)\s*(?:present|current|now|ongoing)',

        # Year to Year format
        r'\b\d{4}\s*(?:-|–|to)\s*\d{4}\b',

        # Year to Present format
        r'\b\d{4}\s*(?:-|–|to)\s*(?:present|current|now|ongoing)\b',

        # Single Month-Year format (start date)
        r'(?:since|from|starting|beginning|commenced)\s+(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s+\d{4}',
        r'(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*\.?\s+\d{4}(?:\s+(?:to|until|through|onwards))',

        # Duration patterns
        r'(?:for|over)\s+(\d+)\s+(?:years|months)',
        r'(\d+)(?:\+)?\s+(?:years|months)(?:\s+of\s+experience)?'
    ]

    location_patterns = [
        # Location with context
        r'(?:located|located in|based in|in|at)\s+([A-Za-z\s,]+)',

        # City, State/Country format
        r'([A-Za-z\s]+),\s*(?:[A-Z]{2}|[A-Za-z]+\s+[A-Za-z]+)',

        # Remote/hybrid work
        r'(remote|hybrid|work from home|wfh|virtual|telecommute)',

        # Location after company
        r'(?:[A-Z][A-Za-z\s&,\-\.]+)(?:\s+in\s+)([A-Za-z\s,]+)'
    ]

    # Expanded job title keywords for better matching
    job_title_keywords = [
        # Technical roles
        "engineer", "developer", "programmer", "architect", "administrator", "technician", "analyst",
        "scientist", "researcher", "designer", "specialist", "consultant", "advisor",

        # Management roles
        "manager", "director", "supervisor", "lead", "head", "chief", "officer", "executive",
        "president", "vp", "vice president", "ceo", "cto", "cfo", "cio", "coo",

        # Other professional roles
        "coordinator", "assistant", "associate", "representative", "agent", "advisor",
        "strategist", "planner", "coach", "trainer", "teacher", "instructor",

        # Entry-level roles
        "intern", "trainee", "apprentice", "junior", "assistant", "entry-level",

        # Seniority indicators
        "senior", "principal", "staff", "distinguished", "executive"
    ]

    # Expanded company keywords
    company_keywords = [
        "company", "corporation", "inc", "incorporated", "llc", "ltd", "limited",
        "group", "associates", "partners", "agency", "firm", "organization",
        "enterprise", "ventures", "solutions", "systems", "technologies", "tech"
    ]

    # First, try to identify distinct job entries by looking for job titles, companies, and dates
    experience_entries = []
    lines = text.split('\n')

    # STRATEGY 1: Group lines into entries based on job titles, companies, and dates
    current_entry = []
    entries = []

    for line in lines:
        line = line.strip()
        if not line:
            if current_entry:
                entries.append('\n'.join(current_entry))
                current_entry = []
            continue

        # Check if this line starts a new entry
        is_new_entry = False

        # Check for job title keywords
        if any(keyword in line.lower() for keyword in job_title_keywords):
            # Only consider it a new entry if it looks like a job title (not just any mention of the keyword)
            if (re.search(r'^[A-Z]', line) or  # Starts with capital letter
                re.search(r'(?:as|position|title|role)\s+(?:a|an)?\s*\w+', line, re.I) or  # Has position indicators
                re.search(r'(?:at|with|for)\s+[A-Z]', line)):  # Has company indicators
                is_new_entry = True

        # Check for company indicators
        if not is_new_entry and any(keyword in line.lower() for keyword in company_keywords):
            if re.search(r'[A-Z][A-Za-z\s&,\-\.]+', line):  # Has capitalized company name
                is_new_entry = True

        # Check for date patterns (often indicate the start of a new entry)
        if not is_new_entry:
            for pattern in date_patterns:
                if re.search(pattern, line, re.I):
                    # Only consider it a new entry if it's at the beginning of the line or after a bullet point
                    if re.search(r'^(\s*[-•*]\s*)?', line):
                        is_new_entry = True
                        break

        if is_new_entry:
            if current_entry:
                entries.append('\n'.join(current_entry))
                current_entry = []
            current_entry.append(line)
        elif current_entry:
            current_entry.append(line)

    # Add the last entry if there is one
    if current_entry:
        entries.append('\n'.join(current_entry))

    # STRATEGY 2: If we couldn't identify distinct entries, try splitting by formatting cues
    if not entries:
        # Try splitting by bullet points, blank lines, or numbering
        entries = re.split(r'\n\s*\n|\n\s*•|\n\s*-|\n\s*\*|\n\s*\d+\.', text)

        # If still no entries, treat the whole text as one entry
        if not entries or (len(entries) == 1 and not entries[0].strip()):
            entries = [text]

    # STRATEGY 3: Process each entry to extract structured information
    for entry in entries:
        if not entry.strip():
            continue

        experience_entry = {
            'title': '',
            'company': '',
            'dates': '',
            'location': '',
            'description': entry.strip(),
            'achievements': [],
            'responsibilities': []
        }

        # Extract job title
        for pattern in title_patterns:
            title_matches = re.findall(pattern, entry, re.I)
            if title_matches:
                # Get the best match (not too short, not too long)
                best_match = ''
                for match in title_matches:
                    if isinstance(match, tuple):
                        match = match[0]  # Some patterns return tuples

                    match = match.strip()
                    if 3 <= len(match) <= 50:  # Reasonable length for a job title
                        # Prefer matches that contain job title keywords
                        if any(keyword in match.lower() for keyword in job_title_keywords):
                            best_match = match
                            break
                        # Otherwise take the first reasonable match
                        if not best_match:
                            best_match = match

                if best_match:
                    experience_entry['title'] = best_match
                    break
                else:
                    experience_entry['title'] = title_matches[0].strip()
                    break

        # If no title found with patterns, look for job title keywords in context
        if not experience_entry['title']:
            for keyword in job_title_keywords:
                if keyword in entry.lower():
                    # Find the sentence containing the keyword
                    sentences = re.split(r'[.!?]\s+', entry)
                    for sentence in sentences:
                        if keyword in sentence.lower():
                            # Extract a reasonable portion of the sentence
                            words = sentence.split()
                            keyword_index = -1
                            for i, word in enumerate(words):
                                if keyword.lower() in word.lower():
                                    keyword_index = i
                                    break

                            if keyword_index >= 0:
                                # Get a window around the keyword (up to 7 words)
                                start_idx = max(0, keyword_index - 3)
                                end_idx = min(len(words), keyword_index + 4)
                                title_phrase = ' '.join(words[start_idx:end_idx])
                                experience_entry['title'] = title_phrase.strip()
                                break
                    if experience_entry['title']:
                        break

        # Extract company
        for pattern in company_patterns:
            company_matches = re.findall(pattern, entry, re.I)
            if company_matches:
                # Get the best match (not too short, not too long)
                best_match = ''
                for match in company_matches:
                    if isinstance(match, tuple):
                        match = match[0]  # Some patterns return tuples

                    match = match.strip()
                    if 2 <= len(match) <= 50:  # Reasonable length for a company name
                        # Prefer matches that contain company keywords
                        if any(keyword in match.lower() for keyword in company_keywords):
                            best_match = match
                            break
                        # Otherwise take the first reasonable match
                        if not best_match:
                            best_match = match

                if best_match:
                    experience_entry['company'] = best_match
                    break
                else:
                    experience_entry['company'] = company_matches[0].strip()
                    break

        # Extract dates
        for pattern in date_patterns:
            date_matches = re.findall(pattern, entry, re.I)
            if date_matches:
                if isinstance(date_matches[0], tuple):
                    # Some patterns return tuples, get the first element
                    experience_entry['dates'] = date_matches[0][0].strip()
                else:
                    experience_entry['dates'] = date_matches[0].strip()
                break

        # Extract location
        for pattern in location_patterns:
            location_matches = re.findall(pattern, entry, re.I)
            if location_matches:
                if isinstance(location_matches[0], tuple):
                    experience_entry['location'] = location_matches[0][0].strip()
                else:
                    experience_entry['location'] = location_matches[0].strip()
                break

        # STRATEGY 4: Extract achievements and responsibilities

        # Look for bullet points (achievements and responsibilities)
        bullet_pattern = r'(?:^|\n)(?:\s*[-•*]\s*|\d+\.\s*)([^\n]+)'
        bullet_points = re.findall(bullet_pattern, entry)

        if bullet_points:
            # Classify each bullet point as achievement or responsibility
            for point in bullet_points:
                point = point.strip()
                if not point:
                    continue

                # Check if it's an achievement (contains action verbs and results)
                is_achievement = False
                achievement_indicators = [
                    r'\b(?:achieved|accomplished|improved|increased|decreased|reduced|enhanced|boosted|grew|expanded|delivered|won|earned|awarded|recognized)\b',
                    r'\b(?:by|resulting in)\s+\d+%',
                    r'\b\d+%\s+(?:increase|decrease|improvement|reduction|growth)',
                    r'\b(?:success|successful|award|recognition|achievement)'
                ]

                for indicator in achievement_indicators:
                    if re.search(indicator, point, re.I):
                        is_achievement = True
                        break

                if is_achievement:
                    experience_entry['achievements'].append(point)
                else:
                    experience_entry['responsibilities'].append(point)

            # If we didn't classify any as achievements or responsibilities, assume they're all responsibilities
            if not experience_entry['achievements'] and not experience_entry['responsibilities']:
                experience_entry['responsibilities'] = [a.strip() for a in bullet_points if a.strip()]

        # Also look for achievement sentences (not in bullet points)
        achievement_sentence_pattern = r'(?:^|\n)(?![-•*\d])([A-Z][^.!?\n]*?(?:achieved|improved|increased|decreased|reduced|enhanced|boosted|grew|expanded|delivered|won|earned|awarded|recognized)[^.!?\n]*?[.!?])'
        achievement_sentences = re.findall(achievement_sentence_pattern, entry, re.I)
        if achievement_sentences:
            experience_entry['achievements'].extend([s.strip() for s in achievement_sentences if s.strip()])

        # STRATEGY 5: Use NER to identify organizations and locations if not found with patterns
        if not experience_entry['company'] or not experience_entry['location']:
            try:
                entry_doc = doc(entry)
                for ent in entry_doc.ents:
                    if ent.label_ == "ORG" and not experience_entry['company']:
                        # Check if it looks like a company (not a university, etc.)
                        if not any(edu_term in ent.text.lower() for edu_term in ["university", "college", "school", "institute"]):
                            experience_entry['company'] = ent.text
                    elif ent.label_ == "GPE" and not experience_entry['location']:
                        experience_entry['location'] = ent.text
            except Exception as e:
                logging.warning(f"Error using NER for experience extraction: {e}")

        # STRATEGY 6: Clean up and validate the extracted information

        # Clean up job title (remove extra whitespace, normalize common titles)
        if experience_entry['title']:
            # Remove any trailing punctuation
            title_text = re.sub(r'[,;.]\s*$', '', experience_entry['title']).strip()

            # Normalize common job title prefixes
            title_text = re.sub(r'^(?:as|position|title|role)(?:\s+of)?(?:\s+a)?(?:\s+the)?\s+', '', title_text, flags=re.I)

            # Clean up whitespace
            title_text = re.sub(r'\s+', ' ', title_text).strip()
            experience_entry['title'] = title_text

        # Clean up company name
        if experience_entry['company']:
            # Remove any trailing punctuation
            company_text = re.sub(r'[,;.]\s*$', '', experience_entry['company']).strip()

            # Normalize common company prefixes
            company_text = re.sub(r'^(?:at|with|for)\s+', '', company_text, flags=re.I)

            # Clean up whitespace
            company_text = re.sub(r'\s+', ' ', company_text).strip()
            experience_entry['company'] = company_text

        # Only add entries that have at least a title or company
        if experience_entry['title'] or experience_entry['company']:
            experience_entries.append(experience_entry)

    # STRATEGY 7: If we still don't have any entries but have text, try a more aggressive approach
    if not experience_entries and text.strip():
        # Look for any mentions of job titles or companies
        for keyword in job_title_keywords:
            if keyword in text.lower():
                # Find the sentence containing the keyword
                sentences = re.split(r'[.!?]\s+', text)
                for sentence in sentences:
                    if keyword in sentence.lower():
                        experience_entry = {
                            'title': sentence.strip(),
                            'company': '',
                            'dates': '',
                            'location': '',
                            'description': text.strip(),
                            'achievements': [],
                            'responsibilities': []
                        }

                        # Try to find a company near this job title
                        for pattern in company_patterns:
                            company_matches = re.findall(pattern, sentence, re.I)
                            if company_matches:
                                experience_entry['company'] = company_matches[0].strip()
                                break

                        experience_entries.append(experience_entry)
                        break
                if experience_entries:
                    break

    logging.info(f"Extracted {len(experience_entries)} experience entries")
    return experience_entries

def create_minimal_profile_from_text(text: str) -> dict:
    """Create a minimal profile from text when full parsing fails."""
    import re
    from datetime import datetime
    from .skill_keywords import ALL_SKILLS

    try:
        # Extract name using regex patterns
        name = ""
        name_patterns = [
            r"(?i)name\s*:\s*([A-Za-z\s\.-]+)",
            r"(?i)^([A-Z][a-z]+\s+[A-Z][a-z]+)$",
            r"(?i)^([A-Z][a-z]+\s+[A-Z]\.?\s+[A-Z][a-z]+)$"
        ]

        lines = [line.strip() for line in text.splitlines() if line.strip()]
        for pattern in name_patterns:
            for line in lines[:10]:  # Check first 10 lines
                match = re.search(pattern, line)
                if match:
                    potential_name = match.group(1).strip()
                    # Verify it looks like a name (1-4 words, properly capitalized)
                    words = potential_name.split()
                    if 1 <= len(words) <= 4 and all(w[0].isupper() for w in words if w and len(w) > 1):
                        name = potential_name
                        break
            if name:
                break

        # Extract email using regex
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        email = emails[0] if emails else ""

        # Extract phone numbers
        phone_patterns = [
            r'\b(?:\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b',  # (123) 456-7890, 123-456-7890
            r'\b\d{3}[-.\s]?\d{4}\b',  # 123-4567
            r'\b\+\d{1,3}[-.\s]?\d{9,10}\b'  # +1 1234567890
        ]

        phone = ""
        for pattern in phone_patterns:
            phones = re.findall(pattern, text)
            if phones:
                phone = phones[0]
                break

        # Extract skills using regex
        skills = []
        for skill in ALL_SKILLS:
            if len(skill) >= 4 and re.search(r'\b' + re.escape(skill) + r'\b', text.lower()):
                skills.append(skill)

        # Limit to top 20 skills
        if len(skills) > 20:
            skills = skills[:20]

        # Extract education (look for degree keywords)
        education = ""
        education_keywords = ["bachelor", "master", "phd", "doctorate", "degree", "diploma", "certificate"]
        for line in lines:
            if any(kw in line.lower() for kw in education_keywords):
                education = line
                break

        # Extract experience (look for job title keywords)
        experience = ""
        job_keywords = ["engineer", "developer", "manager", "analyst", "specialist", "director"]
        for i, line in enumerate(lines):
            if any(kw in line.lower() for kw in job_keywords):
                # Include this line and the next one if available
                experience = line
                if i + 1 < len(lines):
                    experience += "\n" + lines[i + 1]
                break

        return {
            'name': name or 'User',
            'email': email,
            'phone': phone,
            'skills': skills,
            'education': education,
            'experience': experience,
            'raw_text': text[:5000],  # Limit text size
            'uploaded': True,
            'lastUpdated': datetime.now().isoformat()
        }
    except Exception as e:
        logging.error(f"Error in create_minimal_profile_from_text: {e}")
        # Return a very minimal profile if everything fails
        return {
            'name': 'User',
            'email': '',
            'skills': [],
            'education': '',
            'experience': '',
            'uploaded': True,
            'lastUpdated': datetime.now().isoformat()
        }

def extract_target_job(text: str, summary: str, doc) -> str:
    """
    Enhanced target job extraction using multiple strategies:
    1. Look for job titles in the summary/objective section
    2. Look for patterns like "seeking position as X"
    3. Match against comprehensive job title list
    """
    # Common job title patterns
    job_patterns = [
        r'(?:seeking|looking for|interested in|apply for|position as|role as|career as|job as|opportunity as)(?:\s+a)?(?:\s+\w+){0,3}\s+([A-Za-z\s]+(?:developer|engineer|manager|analyst|designer|specialist|consultant|coordinator|director|administrator|officer))',
        r'(?:experienced|senior|junior|lead|principal|chief|head)\s+([A-Za-z\s]+(?:developer|engineer|manager|analyst|designer|specialist|consultant|coordinator|director|administrator|officer))',
        r'(?:^|\n)([A-Za-z\s]{3,30}(?:developer|engineer|manager|analyst|designer|specialist|consultant|coordinator|director|administrator|officer))'
    ]

    target_job = ""

    # Strategy 1: Look for job title patterns in summary
    if summary:
        for pattern in job_patterns:
            matches = re.findall(pattern, summary, re.I)
            if matches:
                # Take the first match
                target_job = matches[0].strip()
                break

    # Strategy 2: If no match in summary, try the full text
    if not target_job:
        for pattern in job_patterns:
            matches = re.findall(pattern, text[:1000], re.I)  # Only search beginning of document
            if matches:
                target_job = matches[0].strip()
                break

    # Strategy 3: Look for exact matches with job titles
    if not target_job:
        for job_title in ALL_JOB_TITLES:
            if job_title in text.lower()[:1000]:
                target_job = job_title
                break

    # Clean up the target job
    if target_job:
        # Remove common prefixes
        target_job = re.sub(r'^(i am a|position:|job title:|role:|career:)\s*', '', target_job, flags=re.I)
        # Limit to a reasonable length
        target_job = ' '.join(target_job.split()[:6])
        # Capitalize properly
        target_job = ' '.join(word.capitalize() if len(word) > 3 else word for word in target_job.split())

    return target_job

def score_section(section_text: str, skills: List[str]) -> Dict[str, float]:
    """Score a section based on various metrics."""
    import textblob

    if not section_text:
        return {
            "completeness": 0.0,
            "clarity": 0.0,
            "impact": 0.0,
            "relevance": 0.0
        }

    # Completeness: length > threshold
    completeness = min(len(section_text) / 200, 1.0)

    # Clarity: TextBlob polarity & subjectivity as proxy
    try:
        blob = textblob.TextBlob(section_text)
        clarity = 1 - abs(blob.sentiment.subjectivity - 0.5)  # closer to 0.5 is more objective
    except:
        clarity = 0.5  # Default if TextBlob fails

    # Impact: count numbers, action verbs, and achievement indicators
    num_numbers = len(re.findall(r"\b\d+\b", section_text))
    action_verbs = ["achieved", "improved", "increased", "decreased", "developed", "created", "implemented", "managed", "led", "designed", "built", "launched", "reduced", "enhanced", "optimized"]
    num_action_verbs = sum(1 for verb in action_verbs if re.search(rf"\b{verb}\b", section_text, re.I))
    impact = min((num_numbers + num_action_verbs) / 5, 1.0)

    # Relevance: keyword overlap with skills
    if skills:
        rel = sum(1 for skill in skills if skill.lower() in section_text.lower()) / len(skills)
    else:
        rel = 0.5  # Default if no skills provided

    return {
        "completeness": round(completeness, 2),
        "clarity": round(clarity, 2),
        "impact": round(impact, 2),
        "relevance": round(rel, 2)
    }

def analyze_language_quality(text: str) -> Dict[str, Any]:
    """Analyze language quality and readability."""
    import textblob

    try:
        blob = textblob.TextBlob(text)

        # Basic metrics
        word_count = len(blob.words)
        sentence_count = len(blob.sentences)
        avg_sentence_length = word_count / max(sentence_count, 1)

        # Sentiment analysis
        polarity = blob.sentiment.polarity
        subjectivity = blob.sentiment.subjectivity

        # Simple readability score (Flesch-Kincaid approximation)
        readability = 206.835 - (1.015 * avg_sentence_length) - (84.6 * (sum(len(word) for word in blob.words) / max(word_count, 1)))
        readability = max(0, min(100, readability))  # Clamp between 0-100

        return {
            "grammar_errors": 0,  # Placeholder; could integrate LanguageTool
            "polarity": round(polarity, 2),
            "subjectivity": round(subjectivity, 2),
            "word_count": word_count,
            "sentence_count": sentence_count,
            "avg_sentence_length": round(avg_sentence_length, 2),
            "readability": round(readability, 2)
        }
    except Exception as e:
        logging.error(f"Error in language quality analysis: {e}")
        return {
            "grammar_errors": 0,
            "polarity": 0,
            "subjectivity": 0,
            "word_count": len(text.split()),
            "sentence_count": len(re.split(r'[.!?]', text)),
            "avg_sentence_length": 0,
            "readability": 0
        }

def generate_ats_report(text: str, email: str, phone: str, skills: List[str],
                        experience: str, education: str, skills_section: str,
                        target_job: str) -> Dict[str, Any]:
    """Generate an ATS (Applicant Tracking System) compatibility report."""
    # Basic checks
    has_email = bool(email)
    has_phone = bool(phone)
    has_skills_section = bool(skills_section)
    has_experience_section = bool(experience)
    has_education_section = bool(education)

    # Check for common ATS issues
    has_tables = "│" in text or "┌" in text or "└" in text or "┘" in text
    has_images = "[image]" in text.lower() or "image:" in text.lower()
    has_complex_formatting = len(re.findall(r'[^\x00-\x7F]', text)) > 10  # Non-ASCII characters

    # Check for keyword density
    if target_job:
        # Get relevant skills for the target job
        relevant_skills = []
        for industry, job_titles in JOB_TITLES.items():
            if any(job.lower() in target_job.lower() for job in job_titles):
                # Get industry-specific skills
                if industry == "Technology":
                    relevant_skills = PROGRAMMING_LANGUAGES + WEB_TECHNOLOGIES + DATA_SCIENCE + CLOUD_DEVOPS
                elif industry == "Business":
                    relevant_skills = BUSINESS_SKILLS + SOFT_SKILLS
                elif industry == "Healthcare":
                    relevant_skills = SOFT_SKILLS + ["healthcare", "medical", "clinical", "patient care"]
                elif industry == "Education":
                    relevant_skills = SOFT_SKILLS + ["curriculum", "teaching", "education", "instruction"]
                break

        # Check which relevant skills are missing
        if relevant_skills:
            skills_lower = [s.lower() for s in skills]
            missing_skills = [s for s in relevant_skills if s.lower() not in skills_lower][:10]  # Limit to 10
        else:
            missing_skills = []
    else:
        missing_skills = []

    # Calculate overall ATS score
    ats_score_components = [
        has_email,
        has_phone,
        has_skills_section,
        has_experience_section,
        has_education_section,
        not has_tables,
        not has_images,
        not has_complex_formatting,
        len(skills) >= 5,
        len(missing_skills) < 5
    ]
    ats_score = sum(1 for c in ats_score_components if c) / len(ats_score_components) * 10  # 0-10 scale

    return {
        "has_email": has_email,
        "has_phone": has_phone,
        "has_skills_section": has_skills_section,
        "has_experience_section": has_experience_section,
        "has_education_section": has_education_section,
        "has_tables": has_tables,
        "has_images": has_images,
        "has_complex_formatting": has_complex_formatting,
        "keyword_match": len(skills) >= 5,
        "keywords_missing": missing_skills,
        "ats_score": round(ats_score, 1)
    }

def check_for_bias(text: str) -> Dict[str, Any]:
    """Check for biased or gendered language in the CV."""
    # Bias terms to check for
    bias_terms = {
        "gendered": ["he", "she", "his", "her", "himself", "herself", "man", "woman", "men", "women"],
        "age_related": ["young", "old", "mature", "senior", "junior", "youthful", "experienced"],
        "problematic": ["blacklist", "whitelist", "master", "slave", "guys", "manpower", "chairman", "manmade", "mankind"]
    }

    found_bias = {}
    for category, terms in bias_terms.items():
        found = [term for term in terms if re.search(rf'\b{term}\b', text, re.I)]
        if found:
            found_bias[category] = found

    # Calculate bias score (0-10, lower is better)
    bias_score = min(10, sum(len(terms) for terms in found_bias.values()))

    return {
        "bias_terms_found": found_bias,
        "bias_score": bias_score,
        "has_bias": bias_score > 0
    }

def generate_recommendations(section_scores: Dict[str, Dict[str, float]],
                            ats_report: Dict[str, Any],
                            bias_report: Dict[str, Any],
                            skills: List[str],
                            target_job: str) -> List[str]:
    """Generate actionable recommendations based on analysis."""
    recommendations = []

    # Section-specific recommendations
    for section, scores in section_scores.items():
        if scores["completeness"] < 0.6:
            recommendations.append(f"Your {section} section could be more detailed.")
        if scores["impact"] < 0.5:
            recommendations.append(f"Try to quantify achievements in your {section} section with numbers and metrics.")
        if scores["clarity"] < 0.5:
            recommendations.append(f"Improve clarity and conciseness in your {section} section.")

    # ATS recommendations
    if not ats_report["has_email"]:
        recommendations.append("Add an email address for better ATS compatibility.")
    if not ats_report["has_phone"]:
        recommendations.append("Include a phone number for better ATS compatibility.")
    if ats_report["has_tables"]:
        recommendations.append("Remove tables from your CV as they can confuse ATS systems.")
    if ats_report["has_images"]:
        recommendations.append("Remove images from your CV as they can confuse ATS systems.")
    if ats_report["has_complex_formatting"]:
        recommendations.append("Simplify formatting to improve ATS compatibility.")

    # Skills recommendations
    if len(skills) < 5:
        recommendations.append("Add more relevant skills to your CV.")
    if ats_report["keywords_missing"] and target_job:
        missing = ", ".join(ats_report["keywords_missing"][:5])
        recommendations.append(f"Consider adding these skills relevant to {target_job}: {missing}")

    # Bias recommendations
    if bias_report["has_bias"]:
        recommendations.append("Remove or rephrase biased or gendered language for a more inclusive CV.")

    # Limit to top 10 recommendations
    return recommendations[:10]

def create_minimal_profile_from_text(text: str) -> Dict[str, Any]:
    """Create a minimal profile when full parsing fails."""
    try:
        # Try to load spaCy model for better extraction
        try:
            nlp = spacy.load("en_core_web_sm")
            doc = nlp(text[:100000])  # Limit to 100K chars to prevent memory issues
        except Exception as e:
            logging.error(f"Error loading spaCy model in minimal profile: {e}")
            doc = None

        # Extract email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        email = emails[0] if emails else ""

        # Extract name (simplified version)
        name = ""
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        for line in lines[:5]:
            if len(line) < 40 and all(w[0].isupper() for w in line.split() if w) and 1 <= len(line.split()) <= 3:
                name = line
                break

        # Extract skills (simplified version)
        skills = []
        for skill in ALL_SKILLS:
            if len(skill) > 3 and re.search(rf'\b{re.escape(skill)}\b', text, re.I):
                skills.append(skill)

        # Extract education (simplified)
        education = ""
        education_keywords = ["bachelor", "master", "phd", "doctorate", "degree", "diploma", "certificate", "university", "college"]
        for line in lines:
            if any(kw in line.lower() for kw in education_keywords):
                education = line
                break

        # Extract experience (simplified)
        experience = ""
        job_keywords = ["engineer", "developer", "manager", "analyst", "specialist", "director"]
        for i, line in enumerate(lines):
            if any(kw in line.lower() for kw in job_keywords):
                experience = line
                if i + 1 < len(lines):
                    experience += "\n" + lines[i + 1]
                break

        # Try to extract sections for summary generation
        sections = {}
        section_headers = {
            "summary": ["summary", "profile", "objective", "professional summary", "about me"],
            "experience": ["experience", "work experience", "professional experience", "employment history"],
            "education": ["education", "academic background", "educational background", "qualifications"],
            "skills": ["skills", "technical skills", "core competencies", "key skills", "expertise"]
        }

        for section_name, headers in section_headers.items():
            for header in headers:
                pattern = re.compile(rf"(?i)\b{re.escape(header)}\b.*?(?:\n|:)(.*?)(?:\n\n|\n[A-Z][A-Za-z\s]+:|\Z)", re.DOTALL)
                match = pattern.search(text)
                if match:
                    sections[section_name] = match.group(1).strip()
                    break

        # Generate a summary using our enhanced function
        summary = ""
        if doc:
            try:
                summary = extract_summary(text, sections, doc)
            except Exception as e:
                logging.error(f"Error generating summary in minimal profile: {e}")

        # If summary generation failed, use a simple approach
        if not summary:
            # Look for an explicit summary section
            if "summary" in sections:
                summary = sections["summary"]
            # Or take the first substantial paragraph
            else:
                for p in lines:
                    if len(p) >= 50 and len(p) <= 500 and '.' in p:
                        summary = p
                        break

        return {
            'name': name or 'User',
            'email': email,
            'skills': skills[:20],  # Limit to 20 skills
            'education': education,
            'experience': experience,
            'summary': summary,
            'raw_text': text[:5000],  # Limit text size
            'uploaded': True
        }
    except Exception as e:
        logging.error(f"Error in create_minimal_profile_from_text: {e}")
        # Return a very minimal profile if everything fails
        return {
            'name': 'User',
            'email': '',
            'skills': [],
            'education': '',
            'experience': '',
            'summary': '',
            'raw_text': text[:1000] if text else '',
            'uploaded': True
        }


# The rest of the class implementation follows...

class AdvancedNLPAnalysis:
    """Advanced NLP analysis for interview responses."""

    def __init__(self, nlp_analyzer: NLPAnalyzer):
        self.nlp_analyzer = nlp_analyzer
        # Example generic answers for originality check
        self.generic_answers = [
            "I am a hard worker and a team player.",
            "My biggest weakness is that I work too hard.",
            "I always strive to do my best in any situation.",
            "I am passionate about learning new things.",
            "I am a perfectionist and pay attention to detail."
        ]
        try:
            from sentence_transformers import SentenceTransformer, util as st_util
            self.st_model = SentenceTransformer('all-MiniLM-L6-v2')
            self.st_util = st_util
        except Exception as e:
            self.st_model = None
            self.st_util = None
            logging.warning(f"SentenceTransformer not available: {e}")

    def analyze_originality(self, response_text: str) -> dict:
        """Check if the answer is generic or potentially copy-pasted."""
        # Use semantic similarity to known generic answers
        try:
            if not self.st_model or not self.st_util:
                return {'originality_score': 1.0, 'is_generic': False, 'most_similar': None}
            # Compute embeddings
            response_emb = self.st_model.encode([response_text], convert_to_tensor=True)
            generic_embs = self.st_model.encode(self.generic_answers, convert_to_tensor=True)
            # Compute similarities
            similarities = self.st_util.cos_sim(response_emb, generic_embs)[0].cpu().numpy()
            max_sim = float(similarities.max())
            most_similar_idx = int(similarities.argmax())
            is_generic = max_sim > 0.75  # threshold for generic/copy-paste
            return {
                'originality_score': 1.0 - max_sim,
                'is_generic': is_generic,
                'most_similar': self.generic_answers[most_similar_idx] if is_generic else None
            }
        except Exception as e:
            logging.error(f"Error in originality analysis: {e}")
            return {'originality_score': 1.0, 'is_generic': False, 'most_similar': None}


    def analyze_emotional_intelligence(self, response_text: str) -> Dict[str, float]:
        """Analyze emotional intelligence in the response."""
        try:
            doc = self.nlp_analyzer.nlp(response_text)

            # Analyze emotional intelligence components
            empathy_score = self._calculate_empathy_score(doc)
            self_awareness_score = self._calculate_self_awareness_score(doc)
            social_skills_score = self._calculate_social_skills_score(doc)

            return {
                'empathy': empathy_score,
                'self_awareness': self_awareness_score,
                'social_skills': social_skills_score
            }
        except Exception as e:
            logging.error(f"Error in emotional intelligence analysis: {e}")
            raise

    def analyze_critical_thinking(self, response_text: str) -> Dict[str, float]:
        """Analyze critical thinking in the response."""
        try:
            doc = self.nlp_analyzer.nlp(response_text)

            # Analyze critical thinking components
            reasoning_score = self._calculate_reasoning_score(doc)
            problem_solving_score = self._calculate_problem_solving_score(doc)
            evidence_based_score = self._calculate_evidence_based_score(doc)

            return {
                'reasoning': reasoning_score,
                'problem_solving': problem_solving_score,
                'evidence_based': evidence_based_score
            }
        except Exception as e:
            logging.error(f"Error in critical thinking analysis: {e}")
            raise

    def analyze_storytelling_ability(self, response_text: str) -> Dict[str, float]:
        """Analyze storytelling ability in the response."""
        try:
            doc = self.nlp_analyzer.nlp(response_text)

            # Analyze storytelling components
            narrative_flow_score = self._calculate_narrative_flow_score(doc)
            character_development_score = self._calculate_character_development_score(doc)
            engagement_score = self._calculate_engagement_score(doc)

            return {
                'narrative_flow': narrative_flow_score,
                'character_development': character_development_score,
                'engagement': engagement_score
            }
        except Exception as e:
            logging.error(f"Error in storytelling analysis: {e}")
            raise

    # Helper methods for emotional intelligence analysis
    def _calculate_empathy_score(self, doc) -> float:
        """Calculate empathy score based on language patterns."""
        # TODO: Implement actual empathy scoring logic
        return 0.8

    def _calculate_self_awareness_score(self, doc) -> float:
        """Calculate self-awareness score based on introspective language."""
        # TODO: Implement actual self-awareness scoring logic
        return 0.75

    def _calculate_social_skills_score(self, doc) -> float:
        """Calculate social skills score based on collaborative language."""
        # TODO: Implement actual social skills scoring logic
        return 0.85

    # Helper methods for critical thinking analysis
    def _calculate_reasoning_score(self, doc) -> float:
        """Calculate reasoning score based on logical language patterns."""
        # TODO: Implement actual reasoning scoring logic
        return 0.9

    def _calculate_problem_solving_score(self, doc) -> float:
        """Calculate problem-solving score based on solution-oriented language."""
        # TODO: Implement actual problem-solving scoring logic
        return 0.85

    def _calculate_evidence_based_score(self, doc) -> float:
        """Calculate evidence-based score based on factual references."""
        # TODO: Implement actual evidence-based scoring logic
        return 0.8

    # Helper methods for storytelling analysis
    def _calculate_narrative_flow_score(self, doc) -> float:
        """Calculate narrative flow score based on story structure."""
        # TODO: Implement actual narrative flow scoring logic
        return 0.9

    def _calculate_character_development_score(self, doc) -> float:
        """Calculate character development score based on role description."""
        # TODO: Implement actual character development scoring logic
        return 0.85

    def _calculate_engagement_score(self, doc) -> float:
        """Calculate engagement score based on interesting language."""
        # TODO: Implement actual engagement scoring logic
        return 0.8

# Export the AdvancedNLPAnalysis class
__all__ = ['AdvancedNLPAnalysis']
